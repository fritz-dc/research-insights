from copy import deepcopy
from pathlib import Path
from datetime import datetime
import hashlib
import json
import os
import re
import time as _time
import unicodedata
from typing import Any, Callable
from urllib.parse import urlencode

import httpx
import numpy as np
import openai as _openai
import pandas as pd
import yaml
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openai import OpenAI
from sklearn.decomposition import NMF
from sklearn.feature_extraction.text import TfidfVectorizer

ROOT = Path(__file__).parent

# ── Config & I/O ──────────────────────────────────────────────────────────────


def load_cfg(path=None):
    """Load params.yaml.

    Resolution order when path is omitted:
    1. {ROOT}/params.yaml
    2. {ROOT}/CONFIG/params.yaml

    This preserves compatibility with both the refactor package layout and the
    earlier repository layout.
    """
    if path is not None:
        cfg_path = Path(path)
    else:
        candidates = [
            ROOT / "params.yaml",
            ROOT / "CONFIG" / "params.yaml",
        ]
        cfg_path = next((p for p in candidates if p.exists()), candidates[0])

    with open(cfg_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def write_json(path: Path, payload: Any) -> Path:
    """Deterministic JSON writer used across manifests and structured outputs."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False, default=str)
    return path


def compute_md5(path: Path, chunk_size: int = 1024 * 1024) -> str | None:
    """Return true content MD5 for a file, or None if the file does not exist."""
    path = Path(path)
    if not path.exists() or not path.is_file():
        return None
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(chunk_size), b""):
            h.update(chunk)
    return h.hexdigest()


def artifact_meta(path: Path, label: str | None = None) -> dict[str, Any]:
    """Build provenance metadata for one artifact path."""
    path = Path(path)
    exists = path.exists()
    return {
        "label": label or path.name,
        "path": str(path),
        "exists": exists,
        "md5": compute_md5(path) if exists and path.is_file() else None,
        "size_bytes": path.stat().st_size if exists else None,
    }


def build_output_path(
    subdir: str,
    fname: str,
    groupby_field: str = None,
    run_date: str = None,
    root: Path = ROOT,
) -> Path:
    """Resolve a canonical OUTPUTS path and create parent directories.

    Without groupby_field/run_date (notebooks 01, 02):
        OUTPUTS/{subdir}/{fname}
    With groupby_field and run_date (legacy grouped outputs):
        OUTPUTS/{groupby_field}/{run_date}/{subdir}/{fname}
    """
    root = Path(root)
    if groupby_field and run_date:
        p = root / "OUTPUTS" / groupby_field / run_date / subdir / fname
    else:
        p = root / "OUTPUTS" / subdir / fname
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


def build_run_output_path(
    subdir: str,
    fname: str,
    groupby_field: str,
    run_date: str,
    run_id: str,
    root: Path = ROOT,
) -> Path:
    """Resolve a run-scoped OUTPUTS path for Notebook 03."""
    root = Path(root)
    p = root / "OUTPUTS" / "runs" / str(groupby_field) / str(run_date) / str(run_id) / subdir / fname
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


def outpath(subdir, fname, root=ROOT, groupby_field=None, run_date=None):
    """Backward-compatible alias for build_output_path()."""
    return build_output_path(
        subdir,
        fname,
        groupby_field=groupby_field,
        run_date=run_date,
        root=root,
    )


def get_run_date() -> str:
    """Return today's date as YYYY-MM-DD for output path nesting."""
    return datetime.now().strftime("%Y-%m-%d")


def canonicalize_filter_spec(filter_logic: str, filters: list[dict[str, Any]] | None) -> dict[str, Any]:
    """Return the canonical filter-spec object used for run hashing/logging."""
    filters = filters or []
    canonical_filters = []
    for item in filters:
        if not isinstance(item, dict):
            raise ValueError(f"Filter entries must be dicts, got {type(item)}")
        canonical_filters.append({k: item[k] for k in sorted(item.keys())})
    return {
        "schema_version": "v1",
        "filter_logic": filter_logic,
        "filters": canonical_filters,
    }


def get_filter_fields_key(filters: list[dict[str, Any]] | None) -> str:
    """Concatenate referenced filter fields, or 'none' when no filters are supplied."""
    filters = filters or []
    fields = sorted({str(f.get("field", "")).strip() for f in filters if str(f.get("field", "")).strip()})
    return "none" if not fields else "__".join(fields)


def get_run_id(groupby_field: str, filter_spec: dict[str, Any] | None = None) -> str:
    """Build a readable run id with scope hash."""
    scope = json.dumps(filter_spec or {"filters": [], "filter_logic": "and"}, sort_keys=True, separators=(",", ":"))
    scope_hash = hashlib.md5(scope.encode("utf-8")).hexdigest()[:8]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{timestamp}_{groupby_field}_{scope_hash}"


def validate_filter_spec(df: pd.DataFrame, filter_logic: str, filters: list[dict[str, Any]] | None) -> None:
    """Validate analysis filter configuration against the dataframe schema."""
    if filter_logic != "and":
        raise ValueError("analysis.filter_logic must be exactly 'and'")
    filters = filters or []
    for rule in filters:
        if "field" not in rule:
            raise ValueError(f"Filter rule missing 'field': {rule}")
        field = rule["field"]
        if field not in df.columns:
            raise ValueError(f"Filter field '{field}' not found in dataframe columns")
        op = rule.get("op")
        if op not in {"eq", "in", "range", "is_null", "not_null"}:
            raise ValueError(f"Unsupported filter op '{op}' in rule: {rule}")
        if op == "eq" and "value" not in rule:
            raise ValueError(f"eq filter missing 'value': {rule}")
        if op == "in":
            values = rule.get("values")
            if not isinstance(values, list) or not values:
                raise ValueError(f"in filter requires non-empty 'values': {rule}")
        if op == "range":
            if "min" not in rule and "max" not in rule:
                raise ValueError(f"range filter requires 'min' and/or 'max': {rule}")


def _coerce_bound_for_series(series: pd.Series, value: Any) -> Any:
    """Coerce range bounds for datetime columns only; forbid other implicit coercions."""
    if pd.api.types.is_datetime64_any_dtype(series):
        return pd.to_datetime(value)
    return value


def apply_filters(
    df: pd.DataFrame,
    filter_logic: str,
    filters: list[dict[str, Any]] | None,
) -> tuple[pd.DataFrame, dict[str, Any]]:
    """Apply validated analysis filters and return the filtered dataframe plus summary."""
    filters = filters or []
    validate_filter_spec(df, filter_logic, filters)
    if not filters:
        return df.copy(), {
            "filter_logic": filter_logic,
            "filters": [],
            "n_rules": 0,
            "input_row_count": int(len(df)),
            "output_row_count": int(len(df)),
            "dropped_row_count": 0,
            "retained_pct": 100.0 if len(df) else 0.0,
            "fields_checked": [],
            "no_rows_after_filter": False,
            "filter_fields_key": "none",
        }

    mask = pd.Series(True, index=df.index)
    for rule in filters:
        field = rule["field"]
        op = rule["op"]
        series = df[field]

        if op == "eq":
            rule_mask = series == rule["value"]
        elif op == "in":
            rule_mask = series.isin(rule["values"])
        elif op == "range":
            rule_mask = pd.Series(True, index=df.index)
            if "min" in rule:
                rule_mask &= series >= _coerce_bound_for_series(series, rule["min"])
            if "max" in rule:
                rule_mask &= series <= _coerce_bound_for_series(series, rule["max"])
        elif op == "is_null":
            rule_mask = series.isna()
        elif op == "not_null":
            rule_mask = series.notna()
        else:
            raise ValueError(f"Unsupported filter op '{op}'")
        mask &= rule_mask.fillna(False)

    out = df.loc[mask].copy()
    retained_pct = round((len(out) / len(df) * 100), 2) if len(df) else 0.0
    return out, {
        "filter_logic": filter_logic,
        "filters": filters,
        "n_rules": len(filters),
        "input_row_count": int(len(df)),
        "output_row_count": int(len(out)),
        "dropped_row_count": int(len(df) - len(out)),
        "retained_pct": retained_pct,
        "fields_checked": [rule["field"] for rule in filters],
        "no_rows_after_filter": out.empty,
        "filter_fields_key": get_filter_fields_key(filters),
    }


def ingest(path):
    """Load SQL extract CSV; parse comma-separated LISTAGG token string into list."""
    raw = pd.read_csv(path)
    date_cols = [c for c in ["posted_date", "funded_date"] if c in raw.columns]
    df = pd.read_csv(path, parse_dates=date_cols) if date_cols else raw
    df["tokens"] = (
        df["tokens"].fillna("").str.split(",").apply(lambda ts: [t.strip() for t in ts if t.strip()])
    )
    return df


def ensure_warning_file(path: Path) -> Path:
    """Create an empty JSONL warnings file if it does not already exist."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        path.write_text("", encoding="utf-8")
    return path


def append_warning(
    path: Path,
    stage_name: str,
    code: str,
    message: str,
    severity: str = "warning",
    context: dict[str, Any] | None = None,
) -> None:
    """Append one structured warning record to a JSONL file."""
    ensure_warning_file(path)
    record = {
        "timestamp": datetime.now().isoformat(),
        "stage_name": stage_name,
        "severity": severity,
        "code": code,
        "message": message,
        "context": context or {},
    }
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False, default=str) + "\n")


def get_warning_count(path: Path) -> int:
    """Count non-empty lines in a JSONL warnings file."""
    path = Path(path)
    if not path.exists():
        return 0
    with open(path, "r", encoding="utf-8") as f:
        return sum(1 for line in f if line.strip())


def get_llm_client() -> OpenAI:
    """Build and return an OpenAI client with SSL verification disabled.

    verify=False is required for the proxy environment. This is intentional
    and not configurable — it reflects a fixed infrastructure constraint.
    Raises ValueError if OPENAI_API_KEY is not set.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("Missing OPENAI_API_KEY environment variable.")
    return OpenAI(api_key=api_key, http_client=httpx.Client(verify=False))


def start_stage_manifest(
    stage_name: str,
    notebook_file: str,
    config_path: str = "params.yaml",
    run_id: str | None = None,
    group_by_field: str | None = None,
    filter_fields_key: str | None = None,
) -> dict[str, Any]:
    """Create the common stage-manifest skeleton shared by all notebooks."""
    config_meta = artifact_meta(ROOT / config_path, label="config")
    started_at = datetime.now().isoformat()
    return {
        "schema_version": "v1",
        "run_id": run_id,
        "group_by_field": group_by_field,
        "filter_fields_key": filter_fields_key,
        "stage_name": stage_name,
        "status": "running",
        "started_at": started_at,
        "completed_at": None,
        "duration_seconds": None,
        "notebook_file": notebook_file,
        "config_path": config_path,
        "config_md5": config_meta["md5"],
        "input_artifacts": [],
        "output_artifacts": [],
        "row_counts": {},
        "key_params": {},
        "warnings_count": 0,
        "warnings_path": None,
    }


def finalize_stage_manifest(
    manifest: dict[str, Any],
    output_path: Path,
    status: str,
    input_artifacts: list[dict[str, Any]] | None = None,
    output_artifacts: list[dict[str, Any]] | None = None,
    row_counts: dict[str, Any] | None = None,
    key_params: dict[str, Any] | None = None,
    warnings_path: Path | None = None,
) -> dict[str, Any]:
    """Finalize and persist a stage manifest."""
    completed_at = datetime.now()
    started_at = datetime.fromisoformat(manifest["started_at"])
    manifest["status"] = status
    manifest["completed_at"] = completed_at.isoformat()
    manifest["duration_seconds"] = round((completed_at - started_at).total_seconds(), 2)
    manifest["input_artifacts"] = input_artifacts or []
    manifest["output_artifacts"] = output_artifacts or []
    manifest["row_counts"] = row_counts or {}
    manifest["key_params"] = key_params or {}
    if warnings_path is not None:
        manifest["warnings_path"] = str(warnings_path)
        manifest["warnings_count"] = get_warning_count(warnings_path)
    write_json(Path(output_path), manifest)
    return manifest


def build_pipeline_manifest(
    output_path: Path,
    run_id: str,
    run_date: str,
    group_by_field: str,
    filter_spec_path: Path,
    filter_summary_path: Path,
    stage_manifest_paths: list[Path],
    warnings_01_path: Path,
    warnings_02_path: Path,
    warnings_03_path: Path,
    final_outputs: dict[str, str],
    config_path: str = "params.yaml",
    filter_fields_key: str | None = None,
    status: str = "success",
) -> dict[str, Any]:
    """Persist the Notebook 03 pipeline manifest."""
    config_meta = artifact_meta(ROOT / config_path, label="config")
    payload = {
        "schema_version": "v1",
        "run_id": run_id,
        "group_by_field": group_by_field,
        "filter_fields_key": filter_fields_key,
        "run_date": run_date,
        "status": status,
        "created_at": datetime.now().isoformat(),
        "config_path": config_path,
        "config_md5": config_meta["md5"],
        "filter_spec_path": str(filter_spec_path),
        "filter_summary_path": str(filter_summary_path),
        "stage_manifests": [str(Path(p)) for p in stage_manifest_paths],
        "warnings_01_path": str(warnings_01_path),
        "warnings_02_path": str(warnings_02_path),
        "warnings_03_path": str(warnings_03_path),
        "warnings_files": [str(warnings_01_path), str(warnings_02_path), str(warnings_03_path)],
        "final_outputs": final_outputs,
    }
    write_json(output_path, payload)
    return payload


# ── Token helpers ─────────────────────────────────────────────────────────────


def tokens_to_str(token_list):
    """Token list → space-joined string for sklearn."""
    return " ".join(token_list) if token_list is not None else ""


def flat_freq(df, col="tokens"):
    """Corpus-wide token frequency Series."""
    return pd.Series([t for ts in df[col] for t in ts]).value_counts()


def token_doc_freq(df):
    """Distinct project count per token — via explode."""
    return (
        df[["project_id", "tokens"]]
        .explode("tokens")
        .rename(columns={"tokens": "token"})
        .drop_duplicates()
        .groupby("token")["project_id"]
        .nunique()
        .rename("doc_count")
    )


# ── Analysis helpers ──────────────────────────────────────────────────────────


def make_vec(min_df, max_df, ngram_range):
    """Shared TF-IDF vectorizer factory."""
    return TfidfVectorizer(
        min_df=min_df,
        max_df=max_df,
        ngram_range=ngram_range,
        token_pattern=r"(?u)\b[a-z][a-z_\-]*\b",
    )


def add_bin(df, bins):
    """Label each project with the first matching analyst-defined bin, else None."""
    df = df.copy()
    df["bin"] = None
    for b in bins:
        mask = (df["posted_date"] >= b["start"]) & (df["posted_date"] <= b["end"])
        df.loc[mask & df["bin"].isna(), "bin"] = b["name"]
    return df


def group_key(keys, group_cols):
    """Normalise groupby key to dict for scalar or tuple keys."""
    return dict(zip(group_cols, keys if isinstance(keys, tuple) else [keys]))


def build_project_topic_bridge(
    weights_df: pd.DataFrame,
    groupby_field: str,
    threshold: float,
) -> pd.DataFrame:
    """Build project-topic bridge using topic_share threshold."""
    required_cols = {groupby_field, "topic_id", "project_id", "weight"}
    missing = required_cols - set(weights_df.columns)
    if missing:
        raise ValueError(f"weights_df missing required columns: {sorted(missing)}")

    totals = (
        weights_df.groupby([groupby_field, "project_id"])["weight"]
        .sum()
        .rename("total_weight")
        .reset_index()
    )

    merged = weights_df.merge(
        totals,
        on=[groupby_field, "project_id"],
        how="left",
        validate="many_to_one",
    )

    merged["topic_share"] = 0.0
    nonzero = merged["total_weight"].fillna(0) > 0
    merged.loc[nonzero, "topic_share"] = (
        merged.loc[nonzero, "weight"] / merged.loc[nonzero, "total_weight"]
    )

    linked = merged[merged["topic_share"] >= threshold].copy()
    linked["topic_key"] = (
        groupby_field
        + "="
        + linked[groupby_field].astype(str)
        + "|topic="
        + linked["topic_id"].astype(str)
    )

    return linked[
        ["topic_key", "project_id", groupby_field, "topic_id", "weight", "topic_share"]
    ]


def slugify_group_value(value: str, max_len: int = 64) -> str:
    """Convert an arbitrary group value to a safe filename component."""
    value = str(value)
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    value = re.sub(r"[^\w]", "_", value)
    value = re.sub(r"_+", "_", value)
    value = value.strip("_")
    return value[:max_len] if value else "unknown"


def load_essay_snippet_lookup(
    project_ids: list[Any],
    data_dir: Path | None = None,
    max_chars: int = 300,
) -> dict[Any, str]:
    """Lazy-load essay text snippets for a project_id set from DATA/project_essay*.csv."""
    data_dir = Path(data_dir) if data_dir is not None else ROOT / "DATA"
    needed = set(project_ids)
    if not needed:
        return {}
    lookup: dict[Any, str] = {}
    essay_files = sorted(data_dir.glob("project_essay*.csv"))
    text_cols = ["essay", "essay_text", "full_text", "project_essay", "text"]
    for fpath in essay_files:
        try:
            cols = pd.read_csv(fpath, nrows=0).columns.tolist()
        except Exception:
            continue
        available_text_col = next((c for c in text_cols if c in cols), None)
        if not available_text_col or "project_id" not in cols:
            continue
        usecols = ["project_id", available_text_col]
        for chunk in pd.read_csv(fpath, usecols=usecols, chunksize=200000):
            sub = chunk[chunk["project_id"].isin(needed - set(lookup.keys()))]
            if sub.empty:
                continue
            for _, row in sub.iterrows():
                text = str(row.get(available_text_col, "") or "")
                text = re.sub(r"\s+", " ", text).strip()
                if text:
                    lookup[row["project_id"]] = text[:max_chars]
            if len(lookup) == len(needed):
                return lookup
    return lookup


def project_insight_for_saved_candidates(insight: dict) -> dict:
    source_topics_out = []
    for src in insight.get("source_topics", []):
        if isinstance(src, dict):
            group = str(src.get("group", "")).strip()
            topic_id = int(float(src.get("topic_id", -1)))
            if group and topic_id >= 0:
                source_topics_out.append(f"{group}|{topic_id}")
        elif isinstance(src, str) and "|" in src:
            source_topics_out.append(src.strip())

    return {
        "title": str(insight.get("title", "")).strip(),
        "what_seeing": str(insight.get("what_seeing", "")).strip(),
        "why_easy_to_miss": str(insight.get("why_easy_to_miss", "")).strip(),
        "source_topics": source_topics_out,
    }


def _norm_text(s):
    s = str(s).lower()
    s = re.sub(r"[^a-z0-9\s]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _token_set(s):
    return set(_norm_text(s).split())

def _jaccard(a, b):
    a = set(a)
    b = set(b)
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)

def _topic_list(val):
    if isinstance(val, list):
        return [str(x).strip() for x in val if str(x).strip()]
    if pd.isna(val):
        return []
    s = str(val).strip()
    if not s:
        return []
    try:
        parsed = json.loads(s)
        if isinstance(parsed, list):
            return [str(x).strip() for x in parsed if str(x).strip()]
    except Exception:
        pass
    return [x.strip() for x in s.split(",") if x.strip()]


def _pair_kind(a, b):
    a_is_key = a["section"] == "key_insights"
    b_is_key = b["section"] == "key_insights"

    if a_is_key and b_is_key:
        return "key_vs_key"

    if (not a_is_key) and (not b_is_key):
        if a["category_bucket"] == b["category_bucket"]:
            return "bg_vs_same_bg"
        return None

    return "bg_vs_key"

def _screen_pair(a, b):
    kind = _pair_kind(a, b)
    if kind is None:
        return None

    topic_overlap = _jaccard(a["verified_topics_list"], b["verified_topics_list"])
    title_overlap = _jaccard(a["title_tokens"], b["title_tokens"])
    text_overlap = _jaccard(a["text_tokens"], b["text_tokens"])

    # One crisp screen for all pair types.
    # We only send plausible overlaps to nano.
    review = (
        topic_overlap >= 0.30
        or title_overlap >= 0.45
        or text_overlap >= 0.40
    )

    if not review:
        return None

    return {
        "pair_kind": kind,
        "topic_overlap": topic_overlap,
        "title_overlap": title_overlap,
        "text_overlap": text_overlap,
    }


# ── Quality report ────────────────────────────────────────────────────────────

HARD_STOPWORDS = {
    "the", "and", "for", "with", "are", "was", "were", "been", "this", "that", "from",
    "they", "them", "their", "will", "would", "should", "could", "have", "has", "had",
    "use", "using", "used", "make", "makes", "making", "get", "gets", "getting",
    "help", "helps", "project", "students", "student", "classroom", "learning",
    "school", "teacher", "teachers", "education", "grade", "grades", "materials",
    "supplies", "tools", "resources", "funded", "funding", "donors", "donor",
}


def quality_report(df, label, doc_freq=None, matrices=None, save_path=None):
    freq = flat_freq(df)
    stops = [t for t in freq.head(200).index if t in HARD_STOPWORDS]
    stats = {
        "checkpoint": label,
        "timestamp": datetime.now().isoformat(),
        "n_projects": len(df),
        "token_count_distribution": df["tokens"].apply(len).describe().to_dict(),
        "vocab": {
            "unique": int(len(freq)),
            "total": int(freq.sum()),
            "stopword_violations": stops,
        },
        "gates": {"no_stopwords": not stops, "violations": stops},
    }
    if doc_freq is not None and len(doc_freq):
        stats["doc_freq"] = {
            "retained": int(len(doc_freq)),
            "min": int(doc_freq.min()),
            "max": int(doc_freq.max()),
            "median": float(doc_freq.median()),
        }
    if matrices:
        stats["matrices"] = {
            k: {
                "shape": list(X.shape),
                "nnz": int(X.nnz),
                "sparsity": round(1 - X.nnz / (X.shape[0] * X.shape[1] or 1), 4),
            }
            for k, X in matrices.items()
        }
    tc = stats["token_count_distribution"]
    print(f"\n{'=' * 55}  [{label}]")
    print(f"  Projects : {stats['n_projects']:,}")
    print(f"  Tok/proj : min={tc['min']:.0f}  p50={tc['50%']:.0f}  max={tc['max']:.0f}")
    print(f"  Vocab    : {stats['vocab']['unique']:,} unique tokens")
    if matrices:
        for k, m in stats["matrices"].items():
            print(f"  {k:20s}: shape={m['shape']}  sparsity={m['sparsity']:.3f}")
    print(f"  Stopwords: {'PASS' if not stops else 'FAIL — ' + str(stops)}")
    print(f"{'=' * 55}\n")
    if save_path:
        write_json(save_path, stats)
    return stats



# ── Notebook 03 orchestration helpers ────────────────────────────────────────
# These helpers were pulled out of 03_insights_generation so the notebook reads
# more like a pipeline and less like a utilities dump. They are intentionally
# parameterized so the notebook can stay explicit about runtime configuration.


def infer_category_family(groupby_field: str) -> str:
    """Classify a grouping field into a coarse sparsity family.

    This powers the notebook's slice-size heuristics without hardcoding logic
    inline in the notebook body.
    """
    field = str(groupby_field).lower()
    if any(tok in field for tok in ["geo", "metro", "county", "city", "state"]):
        return "broad"
    if any(tok in field for tok in ["year", "quarter", "month", "grade", "band"]):
        return "medium"
    return "sparse"


# ── Category TF-IDF + NMF helpers ────────────────────────────────────────────


def cat_tfidf_slice(
    idx: pd.Index,
    df_index: pd.Index,
    X_full,
    feat,
    idf_vals,
    top_n: int,
) -> pd.DataFrame:
    """Score one slice against the rest of the corpus using a shared TF-IDF matrix."""
    # Compare the focal group's rows against everything not in the group.
    rest_idx = df_index.difference(idx)
    X_cat = X_full[idx.tolist()]
    X_rest = X_full[rest_idx.tolist()]

    # Prevalence captures how often a term appears inside the focal slice.
    cat_prev = (X_cat > 0).mean(axis=0).A1
    rest_prev = (X_rest > 0).mean(axis=0).A1 if len(rest_idx) else np.zeros(len(feat))
    tf = X_cat.mean(axis=0).A1

    # Keep the same scoring columns the notebook used before the refactor.
    return pd.DataFrame(
        {
            "token": feat,
            "tf": tf,
            "idf": idf_vals,
            "tfidf": tf * idf_vals,
            "prevalence": cat_prev,
            "contrast": cat_prev - rest_prev,
            "project_count": (X_cat > 0).sum(axis=0).A1.astype(int),
        }
    ).nlargest(top_n, "tfidf")


def choose_n_components(
    n_docs: int,
    retained_vocab: int,
    base_n_components: int,
    slice_rules: dict[str, Any],
) -> int:
    """Choose an NMF topic count using the notebook's existing heuristics."""
    doc_cap = max(4, n_docs // 15)
    vocab_cap = max(4, retained_vocab // 8)
    topic_cap = (
        slice_rules["small_slice_topic_cap"]
        if slice_rules.get("small_slice_mode", False)
        else base_n_components
    )
    return max(4, min(base_n_components, doc_cap, vocab_cap, topic_cap))


def nmf_one(
    docs: list[str],
    ct_cfg: dict[str, Any],
    cn_cfg: dict[str, Any],
    base_n_components: int,
    slice_rules: dict[str, Any],
) -> tuple[pd.DataFrame | None, Any | None, dict[str, Any]]:
    """Fit one NMF slice and return the same triple the notebook used before."""
    vec = make_vec(
        ct_cfg["min_df"],
        ct_cfg["max_df"],
        tuple(ct_cfg.get("ngram_range", [1, 1])),
    )
    X = vec.fit_transform(docs)

    retained_vocab = int(X.shape[1])
    nonzero_tfidf_nnz = int(X.nnz)

    # Keep the early-exit reasons identical so downstream warnings stay stable.
    if retained_vocab < slice_rules["min_retained_vocab"]:
        return None, None, {
            "skip_reason": "low_retained_vocab",
            "retained_vocab": retained_vocab,
            "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
        }

    if nonzero_tfidf_nnz < slice_rules["min_tfidf_nnz"]:
        return None, None, {
            "skip_reason": "low_tfidf_nnz",
            "retained_vocab": retained_vocab,
            "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
        }

    n_components_used = choose_n_components(
        n_docs=len(docs),
        retained_vocab=retained_vocab,
        base_n_components=base_n_components,
        slice_rules=slice_rules,
    )
    if retained_vocab < n_components_used:
        return None, None, {
            "skip_reason": "vocab_below_topic_count",
            "retained_vocab": retained_vocab,
            "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
            "n_components_used": n_components_used,
        }

    model = NMF(
        n_components=n_components_used,
        random_state=cn_cfg["random_seed"],
        init="nndsvd",
        max_iter=cn_cfg["max_iter"],
    )
    W = model.fit_transform(X)
    vocab = vec.get_feature_names_out()

    rows = []
    for i, comp in enumerate(model.components_):
        idx = comp.argsort()[::-1][: cn_cfg["top_words"]]
        rows.append(
            {
                "topic_id": i,
                "top_terms": vocab[idx].tolist(),
                "top_weights": comp[idx].tolist(),
            }
        )

    return pd.DataFrame(rows), W, {
        "n_components_used": n_components_used,
        "retained_vocab": retained_vocab,
        "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
    }


# ── Topic labeling helpers ───────────────────────────────────────────────────


def build_input(
    t_row: pd.Series,
    weights_df: pd.DataFrame,
    pid_text: pd.Series,
    groupby_field: str,
    n_representative: int,
    top_terms_in_prompt: int,
) -> dict[str, Any]:
    """Build one topic-labeling payload from the topic row plus top examples."""
    terms = t_row["top_terms"]
    key_cols = [groupby_field] + (["bin"] if "bin" in t_row.index else [])
    mask = weights_df["topic_id"] == t_row["topic_id"]
    for col in key_cols:
        mask &= weights_df[col] == t_row[col]

    # Representative snippets come from the highest-weighted projects in topic W.
    rep_pids = (
        weights_df[mask]
        .sort_values("weight", ascending=False)["project_id"]
        .tolist()[:n_representative]
    )

    n_uni = top_terms_in_prompt
    n_bi = max(2, top_terms_in_prompt // 2)
    n_nmf = top_terms_in_prompt
    return {
        "group_value": t_row[groupby_field],
        "topic_id": int(t_row["topic_id"]),
        "bin_line": (
            f"\nBin: {t_row['bin']}"
            if "bin" in t_row.index and pd.notna(t_row.get("bin"))
            else ""
        ),
        "unigrams": ", ".join([x for x in terms if " " not in x][:n_uni]),
        "bigrams": ", ".join([x for x in terms if " " in x][:n_bi]),
        "nmf_terms": ", ".join(terms[:n_nmf]),
        "snippets": "\n".join(f"- {pid_text.get(p, '')}" for p in rep_pids),
    }


def _make_label_error(
    inp: dict[str, Any],
    raw_text: str,
    code: str,
    model_labeling: str,
    groupby_field: str,
    error_text: str | None = None,
) -> dict[str, Any]:
    """Return the notebook's standard structured error object for label failures."""
    return {
        "raw": raw_text,
        "parse_error": True,
        "error_code": code,
        "error": error_text,
        "model": model_labeling,
        "timestamp": datetime.now().isoformat(),
        groupby_field: inp["group_value"],
        "topic_id": inp["topic_id"],
    }


def _label_with_retry(
    inp: dict[str, Any],
    *,
    client,
    model_labeling: str,
    system_prompt: str,
    user_prompt_template: str,
    groupby_field: str,
    warnings_path: Path,
    max_retries: int = 3,
    stage_name: str = "03_insights_generation",
) -> dict[str, Any]:
    """Call the labeling model with the notebook's existing retry behavior."""
    text = ""
    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model_labeling,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt_template.format(**inp)},
                ],
            )
            text = resp.choices[0].message.content.strip()
            obj = json.loads(text)

            obj[groupby_field] = str(inp["group_value"])
            obj["topic_id"] = int(inp["topic_id"])
            obj["model"] = model_labeling
            obj["timestamp"] = datetime.now().isoformat()
            return obj
        except json.JSONDecodeError as e:
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
                continue
            append_warning(
                warnings_path,
                stage_name,
                "LABELING_PARSE_FAILURE",
                f"Topic labeling JSON parse failure for {inp['group_value']} / topic {inp['topic_id']}",
                context={
                    "group": inp["group_value"],
                    "topic_id": inp["topic_id"],
                    "error": str(e),
                },
            )
            return _make_label_error(
                inp,
                text,
                "LABELING_PARSE_FAILURE",
                model_labeling=model_labeling,
                groupby_field=groupby_field,
                error_text=str(e),
            )
        except (_openai.RateLimitError, _openai.APITimeoutError) as e:
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
            else:
                append_warning(
                    warnings_path,
                    stage_name,
                    "LABELING_API_FAILURE",
                    f"Topic labeling failed after retries for {inp['group_value']} / topic {inp['topic_id']}",
                    context={
                        "group": inp["group_value"],
                        "topic_id": inp["topic_id"],
                        "error": str(e),
                    },
                )
                return _make_label_error(
                    inp,
                    text or str(e),
                    "LABELING_API_FAILURE",
                    model_labeling=model_labeling,
                    groupby_field=groupby_field,
                    error_text=str(e),
                )
        except Exception as e:
            append_warning(
                warnings_path,
                stage_name,
                "LABELING_API_FAILURE",
                f"Topic labeling failed for {inp['group_value']} / topic {inp['topic_id']}",
                context={
                    "group": inp["group_value"],
                    "topic_id": inp["topic_id"],
                    "error": str(e),
                },
            )
            return _make_label_error(
                inp,
                text or str(e),
                "LABELING_API_FAILURE",
                model_labeling=model_labeling,
                groupby_field=groupby_field,
                error_text=str(e),
            )


def _norm_group_value(value: Any) -> str:
    """Normalize group values before sort-key comparisons."""
    return str(value or "").strip().casefold()


def _safe_topic_id(value: Any) -> int:
    """Best-effort topic-id coercion for sort-key comparisons."""
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return -1


# ── Synthesis helpers ────────────────────────────────────────────────────────


def clean_label(text: Any) -> str:
    """Remove injected token markers before prompt assembly."""
    return re.sub(r"__[a-z_]+__\s*", "", str(text)).strip()


def build_topic_lines(
    df: pd.DataFrame,
    groupby_field: str,
    group: Any | None = None,
    top_terms_count: int = 4,
) -> str:
    """Render labeled topics into the line-oriented prompt format used in Notebook 03."""
    if group is not None:
        df = df[df[groupby_field] == group]

    def _format_top_terms(val: Any, n: int) -> str:
        if n <= 0:
            return ""

        if isinstance(val, list):
            terms = [str(x).strip() for x in val if str(x).strip()]
        elif pd.isna(val):
            terms = []
        else:
            s = str(val).strip()
            if not s:
                terms = []
            else:
                try:
                    parsed = json.loads(s)
                    if isinstance(parsed, list):
                        terms = [str(x).strip() for x in parsed if str(x).strip()]
                    else:
                        terms = [x.strip() for x in s.split(",") if x.strip()]
                except Exception:
                    terms = [x.strip() for x in s.split(",") if x.strip()]

        terms = [clean_label(t) for t in terms[:n]]
        return ", ".join([t for t in terms if t])

    lines = []
    for _, row in df.iterrows():
        top_terms_str = _format_top_terms(row.get("top_terms"), top_terms_count)

        line = (
            f"  {row[groupby_field]} | topic {row.topic_id} | "
            f"label: {clean_label(row.proposed_label)} | "
            f"coherence: {row.coherence_flag} | "
        )

        if top_terms_str:
            line += f"top_terms: {top_terms_str} | "

        line += f"description: {clean_label(row.description)}"
        lines.append(line)

    return "\n".join(lines)


def build_per_group_prompt(
    group: Any,
    group_description: str,
    topic_lines_text: str,
    per_group_instructions: str,
) -> str:
    """Build the per-group synthesis prompt body."""
    return f"""
Below is a list of NMF topics discovered from teacher project request essays on DonorsChoose
for a single group: {group} ({group_description}).
Each topic represents a cluster of essays with similar language, framing, and request patterns.
{topic_lines_text}

{per_group_instructions}
""".strip()


def _call_with_retry(
    prompt: str,
    *,
    client,
    model_name: str,
    system_prompt: str,
    max_retries: int = 3,
) -> str | None:
    """Generic LLM caller with the notebook's existing retry behavior."""
    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
            )
            return resp.choices[0].message.content.strip()
        except (_openai.RateLimitError, _openai.APITimeoutError):
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
            else:
                raise
        except Exception as e:
            print(f"Non-retryable error: {e}")
            return None


def synthesize_one_group(
    group: Any,
    *,
    labels_df: pd.DataFrame,
    groupby_field: str,
    group_description: str,
    per_group_instructions: str,
    client,
    model_name: str,
    system_prompt: str,
    warnings_path: Path,
    outpath_func: Callable[[str, str], Path],
    synthesis_top_terms_count: int = 4,
    max_retries: int = 3,
    stage_name: str = "03_insights_generation",
) -> tuple[Any, str | None]:
    """Run one per-group synthesis pass and persist the raw text output."""
    topic_lines_text = build_topic_lines(
        labels_df,
        groupby_field,
        group=group,
        top_terms_count=synthesis_top_terms_count,
    )
    prompt = build_per_group_prompt(
        group=group,
        group_description=group_description,
        topic_lines_text=topic_lines_text,
        per_group_instructions=per_group_instructions,
    )
    result = _call_with_retry(
        prompt,
        client=client,
        model_name=model_name,
        system_prompt=system_prompt,
        max_retries=max_retries,
    )
    if result is None:
        append_warning(
            warnings_path,
            stage_name,
            "SYNTHESIS_GROUP_FAILED",
            f"Synthesis failed for group '{group}'",
            context={"group": group},
        )
        return group, None

    slug = slugify_group_value(group)
    fpath = outpath_func("analysis", f"llm_synthesis_{slug}.txt")
    with open(fpath, "w", encoding="utf-8") as f:
        f.write(result)
    return group, result


# ── Insight normalization + verification helpers ─────────────────────────────


def strip_json_fences(text: str | None) -> str:
    """Remove optional ```json fences from model output before json.loads()."""
    text = (text or "").strip()
    text = re.sub(r"^\s*```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```\s*$", "", text)
    return text.strip()


def normalize_source_topics(
    source_topics: list[Any] | None,
    required_group_values: list[Any],
) -> list[dict[str, Any]]:
    """Normalize model source_topics into [{'group': ..., 'topic_id': int}] records."""
    out = []

    for src in source_topics or []:
        group = None
        tid = None

        if isinstance(src, dict):
            group = src.get("group", src.get("group_value", ""))
            tid = src.get("topic_id", src.get("topic", src.get("id", "")))
        elif isinstance(src, str):
            s = src.strip()
            if "|" in s:
                left, right = s.rsplit("|", 1)
                group = left.strip()
                tid = right.strip()
            elif re.fullmatch(r"(?i)topic\s+\d+", s):
                tid = re.sub(r"(?i)topic\s+", "", s).strip()

        if group is not None:
            group = str(group).strip()
        if tid is not None:
            tid = str(tid).strip()

        if not tid:
            continue

        try:
            tid = int(float(tid))
        except Exception:
            continue

        if not group or group not in required_group_values:
            continue

        out.append({"group": group, "topic_id": tid})

    # Preserve first-seen ordering while removing duplicate topic references.
    deduped = []
    seen = set()
    for item in out:
        key = (item["group"], item["topic_id"])
        if key not in seen:
            seen.add(key)
            deduped.append(item)

    return deduped


def normalize_insight(
    insight: dict[str, Any],
    required_group_values: list[Any],
) -> dict[str, Any]:
    """Normalize one insight object while preserving claimed source_topics."""
    if not isinstance(insight, dict):
        raise ValueError(f"Insight must be a dict, got {type(insight)}")

    raw_source_topics = deepcopy(insight.get("source_topics", []))
    source_topics = normalize_source_topics(raw_source_topics, required_group_values)

    normalized = {
        "title": str(insight.get("title", "")).strip(),
        "what_seeing": str(insight.get("what_seeing", "")).strip(),
        "why_easy_to_miss": str(insight.get("why_easy_to_miss", "")).strip(),
        "source_topics": source_topics,
        "source_topics_claimed": raw_source_topics,
    }
    if not normalized["title"]:
        raise ValueError(f"Insight missing title: {insight}")
    return normalized


def verify_source_topics(
    insight: dict[str, Any],
    labels_df: pd.DataFrame,
    groupby_field: str,
    required_group_values: list[Any],
    *,
    client,
    model_verify: str,
    system_prompt: str,
    warnings_path: Path,
    max_retries: int = 3,
    stage_name: str = "03_insights_generation",
) -> dict[str, Any]:
    """Verify that each claimed topic directly supports the synthesized insight."""
    title = insight.get("title", "")
    what_seeing = insight.get("what_seeing", "")
    source_topics = insight.get("source_topics", [])
    warning_identifier = (
        title[:120]
        if str(title).strip()
        else str(what_seeing).strip()[:120]
    )

    if not source_topics:
        return insight

    topic_lines = []
    for src in source_topics:
        if isinstance(src, dict):
            group = src.get("group", "")
            tid = str(src.get("topic_id", ""))
        elif isinstance(src, str) and "|" in src:
            group, tid = src.rsplit("|", 1)
        else:
            continue

        match = labels_df[
            (labels_df[groupby_field] == group)
            & (labels_df["topic_id"] == int(float(tid)))
        ]
        if not match.empty:
            row = match.iloc[0]
            topic_lines.append(
                f"  {group}|{tid} | label: {row['proposed_label']} | "
                f"description: {row['description']}"
            )

    if not topic_lines:
        return insight

    prompt = f"""
Insight title: {title}

What we're seeing: {what_seeing}

Claimed source topics:
{chr(10).join(topic_lines)}

Return JSON: {{"verified_topics": [{{"group": "...", "topic_id": <int>}}, ...]}}
""".strip()

    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model_verify,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
                response_format={"type": "json_object"},
            )
            result = json.loads(resp.choices[0].message.content)
            original_n = len(source_topics)
            insight["source_topics"] = normalize_source_topics(
                result.get("verified_topics", source_topics),
                required_group_values,
            )
            verified_n = len(insight.get("source_topics", []))
            if verified_n != original_n:
                print(
                    f"Adjusted source_topics: {title[:80]} | "
                    f"{original_n} -> {verified_n}"
                )
            if verified_n == 0 and original_n > 0:
                print(f"WARNING: all source_topics removed: {title[:80]}")
            return insight
        except (_openai.RateLimitError, _openai.APITimeoutError) as e:
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
            else:
                append_warning(
                    warnings_path,
                    stage_name,
                    "VERIFY_API_FAILURE",
                    f"Verification failed for '{warning_identifier or '[untitled]'}'",
                    context={
                        "title": title or None,
                        "what_seeing_preview": str(what_seeing).strip()[:120] or None,
                        "error": str(e),
                    },
                )
                return insight
        except Exception as e:
            append_warning(
                warnings_path,
                stage_name,
                "VERIFY_API_FAILURE",
                f"Verification failed for '{warning_identifier or '[untitled]'}'",
                context={
                    "title": title or None,
                    "what_seeing_preview": str(what_seeing).strip()[:120] or None,
                    "error": str(e),
                },
            )
            return insight


# ── Evidence + confidence helpers ────────────────────────────────────────────


def get_topic_key(groupby_field: str, group: Any, topic_id: Any) -> str:
    """Build the canonical topic_key used in the bridge table."""
    return f"{groupby_field}={group}|topic={int(float(topic_id))}"


def get_project_ids(
    source_topics: list[Any] | None,
    bridge_lookup: dict[str, pd.DataFrame],
    groupby_field: str,
    max_ids: int | None = None,
) -> tuple[list[Any], dict[str, pd.DataFrame]]:
    """Expand verified source topics into supporting project IDs."""
    if not source_topics:
        return [], {}

    topic_rows_cache: dict[str, pd.DataFrame] = {}
    records = []
    for src in source_topics:
        if isinstance(src, dict):
            group = str(src.get("group", "")).strip()
            topic_id = int(float(src.get("topic_id", -1)))
        elif isinstance(src, str) and "|" in src:
            group, topic_id = src.rsplit("|", 1)
            group = str(group).strip()
            topic_id = int(float(topic_id))
        else:
            continue

        topic_key = get_topic_key(groupby_field, group, topic_id)
        if topic_key not in topic_rows_cache:
            topic_rows_cache[topic_key] = bridge_lookup.get(
                topic_key,
                pd.DataFrame(columns=["project_id", "weight", "topic_share"]),
            )
        records.extend(topic_rows_cache[topic_key]["project_id"].tolist())

    seen = list(dict.fromkeys(records))
    return (seen if max_ids is None else seen[:max_ids]), topic_rows_cache


def iter_candidate_insights(
    data: dict[str, Any],
    output_group_key: str,
):
    """Yield every synthesized insight in the notebook's flat candidate format."""
    idx = 1
    for insight in data.get("key_insights", []):
        yield {
            "insight_id": f"KI_{idx:03d}",
            "section": "key_insights",
            "category_bucket": None,
            "insight": insight,
        }
        idx += 1

    group_idx = 1
    for group_value, items in data.get(output_group_key, {}).items():
        for insight in items:
            yield {
                "insight_id": f"BG_{group_idx:03d}",
                "section": output_group_key,
                "category_bucket": group_value,
                "insight": insight,
            }
            group_idx += 1


def _parse_topic_id(val: Any) -> int:
    """Handle both plain ints and 'Topic N' strings from model output."""
    s = str(val).strip()
    if s.lower().startswith("topic"):
        s = s.split()[-1]
    return int(s)


# ── Tiering + dedupe + report helpers ────────────────────────────────────────


def _verify_insight_list(
    items,
    *,
    labels_df: pd.DataFrame,
    groupby_field: str,
    required_group_values: list[Any],
    client,
    model_verify: str,
    system_prompt: str,
    warnings_path: Path,
    min_source_topics_to_verify: int = 1,
):
    verified_items = []
    changed_count = 0
    dropped_to_zero_count = 0
    topics_before = 0
    topics_after = 0

    for insight in items:
        before_n = len(insight.get("source_topics", []))
        topics_before += before_n

        if before_n >= min_source_topics_to_verify:
            verified = verify_source_topics(
                insight,
                labels_df=labels_df,
                groupby_field=groupby_field,
                required_group_values=required_group_values,
                client=client,
                model_verify=model_verify,
                system_prompt=system_prompt,
                warnings_path=warnings_path,
            )
        else:
            verified = insight

        after_n = len(verified.get("source_topics", []))
        topics_after += after_n

        if after_n != before_n:
            changed_count += 1
        if before_n > 0 and after_n == 0:
            dropped_to_zero_count += 1

        verified_items.append(verified)

    stats = {
        "insight_count": len(items),
        "changed_count": changed_count,
        "dropped_to_zero_count": dropped_to_zero_count,
        "topics_before": topics_before,
        "topics_after": topics_after,
    }
    return verified_items, stats


def _nano_dedup_decision(
    kept_row: dict[str, Any],
    candidate_row: dict[str, Any],
    screen_meta: dict[str, Any],
    *,
    client,
    model_name: str,
    system_prompt: str,
) -> dict[str, Any]:
    """Run the final nano dedupe decision for one screened pair."""
    user_prompt = f"""
Pair type: {screen_meta['pair_kind']}

Insight A (already selected)
section: {kept_row['section']}
category_bucket: {kept_row['category_bucket']}
title: {kept_row['title']}
what_seeing: {kept_row['what_seeing']}
verified_topics: {json.dumps(kept_row['verified_topics_list'], ensure_ascii=False)}

Insight B (candidate)
section: {candidate_row['section']}
category_bucket: {candidate_row['category_bucket']}
title: {candidate_row['title']}
what_seeing: {candidate_row['what_seeing']}
verified_topics: {json.dumps(candidate_row['verified_topics_list'], ensure_ascii=False)}

Similarity hints
topic_overlap: {screen_meta['topic_overlap']:.3f}
title_overlap: {screen_meta['title_overlap']:.3f}
text_overlap: {screen_meta['text_overlap']:.3f}
""".strip()

    resp = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        response_format={"type": "json_object"},
    )
    return json.loads(resp.choices[0].message.content)


# ── DOCX report helpers ───────────────────────────────────────────────────────


def build_looker_project_url(
    *,
    base_url: str,
    project_ids: list[Any],
    filter_field: str,
    fields: list[str] | str,
    limit: int = 500,
    max_ids: int = 100,
) -> str:
    """Build a Looker Explore URL for a capped list of supporting project IDs.

    Behavior:
    - Returns "" when project_ids is empty after normalization.
    - Joins `fields` with commas before URL encoding.
    - Encodes the filter as f[<filter_field>]=comma-separated-project-ids.
    - Deduplicates project_ids while preserving their ranked order.
    - Caps IDs at max_ids before building the URL.
    """
    def _normalize_project_id(value: Any) -> str:
        if pd.isna(value):
            return ""
        try:
            numeric = float(value)
            if numeric.is_integer():
                return str(int(numeric))
        except (TypeError, ValueError):
            pass
        return str(value).strip()

    normalized_ids = []
    for value in project_ids:
        pid = _normalize_project_id(value)
        if pid:
            normalized_ids.append(pid)

    normalized_ids = list(dict.fromkeys(normalized_ids))[:max_ids]
    if not normalized_ids:
        return ""

    if isinstance(fields, (list, tuple)):
        fields_param = ",".join(str(f).strip() for f in fields if str(f).strip())
    else:
        fields_param = str(fields).strip()

    if not fields_param:
        raise ValueError("build_looker_project_url requires at least one field")

    params = {
        "fields": fields_param,
        f"f[{filter_field}]": ",".join(normalized_ids),
        "limit": str(limit),
    }
    return f"{base_url}?{urlencode(params)}"


def add_hyperlink(paragraph, text: str, url: str):
    """Add an external hyperlink to a python-docx paragraph."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p.runs[0].font.name = "Arial"
    return p


def add_insight(
    doc,
    insight: dict[str, Any],
    *,
    include_looker_link: bool = True,
    looker_link_text: str = "Project essays for this insight",
) -> None:
    """Add one accepted insight to the final DOCX report."""
    p = doc.add_paragraph()
    run = p.add_run(insight.get("title", ""))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3e, 0x00, 0xc9)

    # Keep the body labels and font sizing identical to the original notebook.
    for label, key in [
        ("What we're seeing:", "what_seeing"),
        ("Why this is easy to miss:", "why_easy_to_miss"),
    ]:
        p = doc.add_paragraph()
        label_run = p.add_run(label + "  ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        body_run = p.add_run(insight.get(key, ""))
        body_run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    if include_looker_link and insight.get("looker_url"):
        p = doc.add_paragraph()
        label_run = p.add_run("Explore supporting projects: ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        add_hyperlink(p, looker_link_text, insight["looker_url"])
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


def add_insight_meta_line(doc, insight: dict, font_size_pt: float = 8.5) -> None:
    """Add a compact italic meta line above an insight body."""
    mean_topic_fit = insight.get("mean_topic_share_all_verified_topics", None)
    if isinstance(mean_topic_fit, (int, float)):
        mean_topic_fit = f"{round(mean_topic_fit * 100):.0f}%"
    else:
        mean_topic_fit = "—"

    text = (
        f"Supporting projects: {insight.get('supporting_project_count', '—')}  |  "
        f"Verified source topics: {insight.get('verified_topic_count', '—')}  |  "
        f"Average topic fit: {mean_topic_fit}"
    )
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(font_size_pt)
    p.paragraph_format.space_after = Pt(2)


def _add_report_summary(
    doc, structured, output_group_key,
    key_tiers, other_tiers, key_label, other_label,
    project_count=None,
    run_id=None,
) -> None:
    all_items = structured.get("key_insights", []) + [
        i for items in structured.get(output_group_key, {}).values() for i in items
    ]
    key_n   = sum(1 for i in all_items if str(i.get("tier", "")).lower() in key_tiers)
    other_n = sum(1 for i in all_items if str(i.get("tier", "")).lower() in other_tiers)
    parts = []
    if project_count is not None:
        parts.append(f"Projects in study: {project_count:,}")
    parts += [f"{key_label}: {key_n}", f"{other_label}: {other_n}"]
    if run_id is not None:
        date_time = "_".join(str(run_id).split("_")[:2])
        parts.append(f"Run ID: {date_time}")
    p = doc.add_paragraph()
    run = p.add_run("  |  ".join(parts))
    run.italic = True
    run.font.size = Pt(9)
    doc.add_paragraph()


def build_packaged_report_docx(
    *,
    structured: dict,
    output_path,
    output_group_key: str = "by_group",
    report_cfg: dict,
    project_count: int | None = None,
    run_id: str | None = None,
    normal_font_name: str = "Arial",
    normal_font_size_pt: float = 10.0,
    margin_inches: float = 1.0,
) -> None:
    from docx import Document
    from docx.shared import Inches

    title = report_cfg.get("report_title", "Trend Report")
    main_label = report_cfg.get("report_main_section_label", "Main Insights")
    appendix_label = report_cfg.get("report_appendix_section_label", "Appendix")
    main_cross_label = report_cfg.get("report_main_cross_label", "Cross-Category Similarities")
    main_by_group_label = report_cfg.get("report_main_by_group_label", "Group-Specific Findings")
    appendix_cross_label = report_cfg.get("report_appendix_cross_label", "Additional Cross-Category Insights")
    appendix_by_group_label = report_cfg.get("report_appendix_by_group_label", "Additional Group-Specific Findings")
    incl_meta = report_cfg.get("report_include_meta_line", True)
    incl_summary = report_cfg.get("report_include_signal_summary", True)
    incl_looker_link = report_cfg.get("report_include_looker_link", True)
    looker_link_text = report_cfg.get("report_looker_link_text", "Project essays for this insight")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(margin_inches)
    style = doc.styles["Normal"]
    style.font.name = normal_font_name
    style.font.size = Pt(normal_font_size_pt)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x3e, 0x00, 0xc9)
    p.paragraph_format.space_after = Pt(6)

    if incl_summary:
        all_items = structured.get("key_insights", []) + [
            i for items in structured.get(output_group_key, {}).values() for i in items
        ]
        parts = []
        if project_count is not None:
            parts.append(f"Projects in study: {project_count:,}")
        parts.append(f"Accepted insights: {len(all_items)}")
        if run_id is not None:
            parts.append(f"Run ID: {'_'.join(str(run_id).split('_')[:2])}")
        p = doc.add_paragraph()
        run = p.add_run("  |  ".join(parts))
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()

    def _items(section_name):
        cross = [i for i in structured.get("key_insights", []) if i.get("report_section") == section_name]
        by_group = {
            g: [i for i in items if i.get("report_section") == section_name]
            for g, items in structured.get(output_group_key, {}).items()
        }
        return cross, by_group

    def _render_bucket(section_title, cross_title, by_group_title, cross_key, by_group_key):
        cross, by_group = _items(cross_key)
        cross2, by_group2 = _items(by_group_key)
        cross = cross + cross2
        merged_by_group = {}
        for g, items in by_group.items():
            merged_by_group[g] = merged_by_group.get(g, []) + items
        for g, items in by_group2.items():
            merged_by_group[g] = merged_by_group.get(g, []) + items

        if not cross and not any(merged_by_group.values()):
            return

        p = doc.add_paragraph()
        run = p.add_run(section_title)
        run.bold = True
        run.underline = True
        run.font.name = "Arial"
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        if cross:
            p = doc.add_paragraph()
            run = p.add_run(cross_title)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(10)
            for insight in sorted(cross, key=lambda x: (x.get("report_order", 9999), x.get("title", ""))):
                if incl_meta:
                    add_insight_meta_line(doc, insight)
                add_insight(
                    doc,
                    insight,
                    include_looker_link=incl_looker_link,
                    looker_link_text=looker_link_text,
                )

        if any(merged_by_group.values()):
            p = doc.add_paragraph()
            run = p.add_run(by_group_title)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(10)
            for group, items in sorted(merged_by_group.items()):
                if not items:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(str(group))
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                for insight in sorted(items, key=lambda x: (x.get("report_order", 9999), x.get("title", ""))):
                    if incl_meta:
                        add_insight_meta_line(doc, insight)
                    add_insight(
                        doc,
                        insight,
                        include_looker_link=incl_looker_link,
                        looker_link_text=looker_link_text,
                    )

    _render_bucket(main_label, main_cross_label, main_by_group_label, "main_cross", "main_by_group")
    _render_bucket(appendix_label, appendix_cross_label, appendix_by_group_label, "appendix_cross", "appendix_by_group")
    doc.save(output_path)


def build_bridge_lookup(project_topic_bridge_df: pd.DataFrame) -> dict[str, pd.DataFrame]:
    return {
        topic_key: grp[["project_id", "weight", "topic_share"]].copy()
        for topic_key, grp in project_topic_bridge_df.groupby("topic_key", observed=True)
    }


def build_label_index(
    labels_df: pd.DataFrame,
    groupby_field: str,
    warnings_path: Path | None = None,
    stage_name: str = "03_insights_generation",
) -> dict[tuple[str, int], Any]:
    out = {}
    for _, row in labels_df.iterrows():
        key = (str(row[groupby_field]), int(row["topic_id"]))
        if key in out and warnings_path is not None:
            append_warning(
                warnings_path,
                stage_name,
                "DUPLICATE_TOPIC_LABEL_ROW",
                f"Duplicate labels_df row for {key}; keeping first occurrence",
                context={"group": key[0], "topic_id": key[1]},
            )
            continue
        out[key] = row
    return out


def summarize_insight_support(
    candidate: dict[str, Any],
    *,
    groupby_field: str,
    run_id: str,
    bridge_lookup: dict[str, pd.DataFrame],
    label_index: dict[tuple[str, int], Any],
    top_project_id_limit: int = 100,
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    insight = candidate["insight"]
    source_topics = insight.get("source_topics", [])
    claimed_topic_count = len(insight.get("source_topics_claimed", []))

    support_rows = []
    ranking_frames = []

    for src in source_topics:
        if not isinstance(src, dict):
            continue

        group = str(src.get("group", "")).strip()
        topic_id = int(float(src.get("topic_id", -1)))
        topic_key = get_topic_key(groupby_field, group, topic_id)
        topic_rows = bridge_lookup.get(
            topic_key,
            pd.DataFrame(columns=["project_id", "weight", "topic_share"]),
        )
        label_row = label_index.get((group, topic_id))

        if not topic_rows.empty:
            ranking_frames.append(
                topic_rows[["project_id", "topic_share", "weight"]].copy()
            )

        support_rows.append(
            {
                "run_id": run_id,
                "insight_id": candidate["insight_id"],
                "section": candidate["section"],
                "category_bucket": candidate["category_bucket"],
                "group_by_field": groupby_field,
                "group_value": group,
                "topic_id": topic_id,
                "topic_label": label_row["proposed_label"] if label_row is not None else "[not found in labels_df]",
                "topic_description": label_row["description"] if label_row is not None else "[not found in labels_df]",
                "coherence_flag": label_row["coherence_flag"] if label_row is not None else "unknown",
                "supporting_project_count": int(topic_rows["project_id"].nunique()) if not topic_rows.empty else 0,
                "mean_topic_share": float(topic_rows["topic_share"].mean()) if not topic_rows.empty else 0.0,
                "median_topic_share": float(topic_rows["topic_share"].median()) if not topic_rows.empty else 0.0,
            }
        )

    if ranking_frames:
        project_scores = (
            pd.concat(ranking_frames, ignore_index=True)
            .groupby("project_id", as_index=False)
            .agg(
                total_topic_share=("topic_share", "sum"),
                total_weight=("weight", "sum"),
            )
        )
        project_scores["project_id_numeric"] = pd.to_numeric(
            project_scores["project_id"],
            errors="coerce",
        )
        project_scores = project_scores.sort_values(
            ["total_topic_share", "total_weight", "project_id_numeric", "project_id"],
            ascending=[False, False, True, True],
            na_position="last",
        )
        top_project_ids = project_scores["project_id"].tolist()[:top_project_id_limit]
        supporting_project_count = int(len(project_scores))
    else:
        top_project_ids = []
        supporting_project_count = 0

    verified_topic_count = len(support_rows)
    verification_ratio = (
        verified_topic_count / claimed_topic_count if claimed_topic_count > 0 else 0.0
    )
    mean_topic_share_all_verified_topics = (
        float(np.mean([r["mean_topic_share"] for r in support_rows])) if support_rows else 0.0
    )

    flat_row = {
        "run_id": run_id,
        "insight_id": candidate["insight_id"],
        "section": candidate["section"],
        "category_bucket": candidate["category_bucket"],
        "title": insight.get("title", ""),
        "what_seeing": insight.get("what_seeing", ""),
        "why_easy_to_miss": insight.get("why_easy_to_miss", ""),
        "source_topics_verified": source_topics,
        "source_topics_claimed": insight.get("source_topics_claimed", []),
        "claimed_topic_count": int(claimed_topic_count),
        "verified_topic_count": int(verified_topic_count),
        "verification_ratio": float(verification_ratio),
        "supporting_project_count": int(supporting_project_count),
        "mean_topic_share_all_verified_topics": float(mean_topic_share_all_verified_topics),
        "top_project_ids": top_project_ids,
    }
    return flat_row, support_rows


def build_verified_insight_tables(
    insights_data: dict[str, Any],
    output_group_key: str,
    *,
    groupby_field: str,
    bridge_lookup: dict[str, pd.DataFrame],
    label_index: dict[tuple[str, int], Any],
    run_id: str | None = None,
    top_project_id_limit: int = 100,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    flat_rows = []
    support_rows = []

    for candidate in iter_candidate_insights(insights_data, output_group_key):
        flat_row, topic_rows = summarize_insight_support(
            candidate,
            groupby_field=groupby_field,
            run_id=run_id,
            bridge_lookup=bridge_lookup,
            label_index=label_index,
            top_project_id_limit=top_project_id_limit,
        )
        flat_rows.append(flat_row)
        support_rows.extend(topic_rows)

    return pd.DataFrame(flat_rows), pd.DataFrame(support_rows)


def apply_deterministic_packaging(
    insights_flat_df: pd.DataFrame,
    *,
    output_group_key: str,
    packaging_cfg: dict[str, Any],
) -> dict[str, pd.DataFrame]:
    pack_df = insights_flat_df.copy()
    pack_df["verified_topic_count"] = pack_df["verified_topic_count"].fillna(0).astype(int)
    pack_df["claimed_topic_count"] = pack_df["claimed_topic_count"].fillna(0).astype(int)
    pack_df["supporting_project_count"] = pack_df["supporting_project_count"].fillna(0).astype(int)
    pack_df["mean_topic_share_all_verified_topics"] = (
        pack_df["mean_topic_share_all_verified_topics"].fillna(0.0).astype(float)
    )
    if "verification_ratio" not in pack_df.columns:
        pack_df["verification_ratio"] = np.where(
            pack_df["claimed_topic_count"] > 0,
            pack_df["verified_topic_count"] / pack_df["claimed_topic_count"],
            0.0,
        )

    accepted_df = pack_df[
        (pack_df["verified_topic_count"] >= packaging_cfg["min_verified_topic_count"])
        & (pack_df["supporting_project_count"] >= packaging_cfg["min_supporting_project_count"])
        & (pack_df["mean_topic_share_all_verified_topics"] >= packaging_cfg["min_mean_topic_share"])
    ].copy()

    accepted_df = accepted_df.sort_values(
        ["supporting_project_count", "verified_topic_count", "mean_topic_share_all_verified_topics", "title"],
        ascending=[False, False, False, True],
    ).reset_index(drop=True)

    empty_df = accepted_df.iloc[0:0].copy()
    return {
        "accepted_df": accepted_df,
        "main_cross_df": empty_df.copy(),
        "main_by_group_df": empty_df.copy(),
        "appendix_df": accepted_df.copy(),
    }


def dedupe_packaged_insights(
    accepted_df: pd.DataFrame,
    *,
    dedupe_cfg: dict[str, Any],
) -> tuple[pd.DataFrame, pd.DataFrame]:
    working_df = accepted_df.copy()
    working_df["verified_topics_list"] = working_df["source_topics_verified"].apply(_topic_list)
    working_df["title_tokens"] = working_df["title"].apply(_token_set)
    working_df["text_tokens"] = (
        working_df["title"].fillna("") + " " + working_df["what_seeing"].fillna("")
    ).apply(_token_set)

    working_df = working_df.sort_values(
        ["mean_topic_share_all_verified_topics", "supporting_project_count", "verification_ratio"],
        ascending=False,
    ).reset_index(drop=True)

    kept_rows = []
    audit_rows = []

    for _, row in working_df.iterrows():
        row_dict = row.to_dict()
        drop_current = False

        for kept in kept_rows:
            screen_meta = _screen_pair(row_dict, kept)
            if screen_meta is None or screen_meta["pair_kind"] == "bg_vs_key":
                continue

            topic_overlap = screen_meta["topic_overlap"]
            title_overlap = screen_meta["title_overlap"]
            text_overlap = screen_meta["text_overlap"]

            is_obvious_duplicate = (
                (topic_overlap >= dedupe_cfg["topic_overlap_high"])
                or (
                    topic_overlap >= dedupe_cfg["topic_overlap_min"]
                    and title_overlap >= dedupe_cfg["title_overlap_min"]
                    and text_overlap >= dedupe_cfg["text_overlap_min"]
                )
            )

            if is_obvious_duplicate:
                drop_current = True
                audit_rows.append(
                    {
                        "dropped_insight_id": row_dict["insight_id"],
                        "matched_kept_insight_id": kept["insight_id"],
                        "pair_kind": screen_meta["pair_kind"],
                        "topic_overlap": topic_overlap,
                        "title_overlap": title_overlap,
                        "text_overlap": text_overlap,
                        "reason": "deterministic_obvious_duplicate",
                    }
                )
                break

        if not drop_current:
            kept_rows.append(row_dict)

    return pd.DataFrame(kept_rows).copy(), pd.DataFrame(audit_rows).copy()


def assign_topline_sections_simple(
    curated_df: pd.DataFrame,
    *,
    output_group_key: str,
    main_cross_limit: int,
) -> pd.DataFrame:
    rank_cols = [
        "supporting_project_count",
        "verified_topic_count",
        "mean_topic_share_all_verified_topics",
        "title",
    ]
    rank_asc = [False, False, False, True]

    working_df = curated_df.copy()

    main_cross_df = (
        working_df[working_df["section"] == "key_insights"]
        .sort_values(rank_cols, ascending=rank_asc)
        .head(main_cross_limit)
        .copy()
    )
    main_cross_df["report_section"] = "main_cross"
    main_cross_df["report_order"] = range(1, len(main_cross_df) + 1)

    main_by_group_df = (
        working_df[working_df["section"] == output_group_key]
        .sort_values(
            ["category_bucket", "supporting_project_count", "verified_topic_count", "mean_topic_share_all_verified_topics", "title"],
            ascending=[True, False, False, False, True],
        )
        .groupby("category_bucket", as_index=False, sort=True)
        .head(1)
        .copy()
    )
    main_by_group_df["report_section"] = "main_by_group"
    main_by_group_df["report_order"] = range(1, len(main_by_group_df) + 1)

    main_ids = set(main_cross_df["insight_id"]).union(set(main_by_group_df["insight_id"]))
    remainder_df = working_df[~working_df["insight_id"].isin(main_ids)].copy()

    appendix_cross_df = (
        remainder_df[remainder_df["section"] == "key_insights"]
        .sort_values(rank_cols, ascending=rank_asc)
        .copy()
    )
    appendix_cross_df["report_section"] = "appendix_cross"
    appendix_cross_df["report_order"] = range(1, len(appendix_cross_df) + 1)

    appendix_by_group_df = (
        remainder_df[remainder_df["section"] == output_group_key]
        .sort_values(
            ["category_bucket", "supporting_project_count", "verified_topic_count", "mean_topic_share_all_verified_topics", "title"],
            ascending=[True, False, False, False, True],
        )
        .copy()
    )
    appendix_by_group_df["report_section"] = "appendix_by_group"
    appendix_by_group_df["report_order"] = (
        appendix_by_group_df.groupby("category_bucket", sort=True).cumcount() + 1
    )

    return pd.concat(
        [main_cross_df, main_by_group_df, appendix_cross_df, appendix_by_group_df],
        ignore_index=True,
        sort=False,
    )


def build_structured_from_curated(
    curated_df: pd.DataFrame,
    *,
    output_group_key: str,
) -> dict[str, Any]:
    structured = {"key_insights": [], output_group_key: {}}

    for _, row in curated_df.sort_values(["report_section", "report_order", "title"]).iterrows():
        looker_url = row.get("looker_url", "")
        if pd.isna(looker_url):
            looker_url = ""

        top_project_ids = row.get("top_project_ids", [])
        if pd.isna(top_project_ids) if not isinstance(top_project_ids, list) else False:
            top_project_ids = []
        if not isinstance(top_project_ids, list):
            top_project_ids = []

        item = {
            "insight_id": row["insight_id"],
            "title": row["title"],
            "what_seeing": row["what_seeing"],
            "why_easy_to_miss": row["why_easy_to_miss"],
            "source_topics": row["source_topics_verified"],
            "supporting_project_count": int(row["supporting_project_count"]),
            "verified_topic_count": int(row["verified_topic_count"]),
            "verification_ratio": float(row["verification_ratio"]),
            "mean_topic_share_all_verified_topics": float(row["mean_topic_share_all_verified_topics"]),
            "top_project_ids": top_project_ids,
            "looker_url": looker_url,
            "report_section": row["report_section"],
            "report_order": int(row["report_order"]),
            "warnings": [],
        }
        if row["section"] == "key_insights":
            structured["key_insights"].append(item)
        else:
            structured[output_group_key].setdefault(row["category_bucket"], []).append(item)

    return structured


__all__ = [
    "ROOT",
    "load_cfg",
    "write_json",
    "compute_md5",
    "artifact_meta",
    "build_output_path",
    "build_run_output_path",
    "outpath",
    "get_run_date",
    "canonicalize_filter_spec",
    "get_filter_fields_key",
    "get_run_id",
    "validate_filter_spec",
    "apply_filters",
    "ingest",
    "ensure_warning_file",
    "append_warning",
    "get_warning_count",
    "get_llm_client",
    "start_stage_manifest",
    "finalize_stage_manifest",
    "build_pipeline_manifest",
    "tokens_to_str",
    "flat_freq",
    "token_doc_freq",
    "make_vec",
    "add_bin",
    "group_key",
    "build_project_topic_bridge",
    "slugify_group_value",
    "load_essay_snippet_lookup",
    "quality_report",
    "_norm_text",
    "_token_set",
    "_jaccard",
    "_topic_list",
    "_stronger_key",
    "_pair_kind",
    "_screen_pair",
    "cat_tfidf_slice",
    "choose_n_components",
    "nmf_one",
    "build_input",
    "_make_label_error",
    "_label_with_retry",
    "_norm_group_value",
    "_safe_topic_id",
    "clean_label",
    "build_topic_lines",
    "build_per_group_prompt",
    "_call_with_retry",
    "synthesize_one_group",
    "strip_json_fences",
    "normalize_source_topics",
    "normalize_insight",
    "verify_source_topics",
    "get_topic_key",
    "iter_candidate_insights",
    "_parse_topic_id",
    "build_bridge_lookup",
    "build_label_index",
    "summarize_insight_support",
    "build_verified_insight_tables",
    "apply_deterministic_packaging",
    "dedupe_packaged_insights",
    "assign_topline_sections_simple",
    "build_structured_from_curated",
    "add_heading",
    "add_insight",
    "add_insight_meta_line",
    "build_packaged_report_docx",
    "project_insight_for_saved_candidates",
    "_verify_insight_list"
]
