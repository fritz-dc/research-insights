"""
utils.py — Trend Tracker shared utilities  (pipeline v1.0.0)

All functions used across the three pipeline notebooks live here. Notebooks
should contain only cell-level orchestration, parameter setup, and display
logic — no reusable logic of their own.

Section map (in order):
    1.  Imports & module-level constants
    2.  Config & I/O
    3.  Run identity & filter helpers
    4.  Stage & pipeline manifests
    5.  Warning file helpers
    6.  LLM client
    7.  Ingest
    8.  Token helpers
    9.  Consolidation helpers          (NB01 Steps 3-4)
    10. Analysis helpers
    11. Quality
    12. TF-IDF & NMF helpers           (NB03 Steps 1-4)
    13. Enrichment helpers             (NB02 Passes A-C)
    14. Topic labeling helpers         (NB03 Step 5)
    15. Synthesis helpers              (NB03 Step 5)
    16. Insight normalization & verification
    17. Dedup helpers
    18. Evidence & support tables
    19. Packaging & tiering
    20. DOCX report helpers
    21. __all__
"""

# ── 1. Imports & module-level constants ───────────────────────────────────────

from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urlencode
import hashlib
import json
import os
import re
import time as _time
import unicodedata

import httpx
import numpy as np
import openai as _openai
import pandas as pd
import simplemma
import yaml
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openai import OpenAI
from sklearn.decomposition import NMF
from sklearn.feature_extraction.text import TfidfVectorizer

# Root of the project tree; used to resolve all relative paths.
ROOT = Path(__file__).parent

# Pipeline version string; recorded in every stage manifest and the pipeline
# manifest so that any output artifact can be traced back to a specific release.
PIPELINE_VERSION = "1.0.0"

# Regex that flags likely simplemma truncation artifacts (stems ending in
# common consonant clusters where a vowel is expected). Used by
# build_consolidation_candidates() to mark flagged tokens for review.
_TRUNC_RE = re.compile(r"(at|iv|iz|az|ig|if|ic|olog)$")


# ── 2. Config & I/O ───────────────────────────────────────────────────────────


def _version_sort_key(path: Path) -> tuple:
    """Sort versioned files so higher explicit versions win; plain names sort last."""
    m = re.search(r"v(\d+)", path.stem)
    return (int(m.group(1)) if m else -1, path.name.lower())


def _resolve_params_path(base_dirs: list[Path]) -> Path | None:
    """Return the highest-versioned params*.yaml / params*.yml across base_dirs."""
    matches: list[Path] = []
    seen: set[str] = set()
    for base in base_dirs:
        for pattern in ("params*.yaml", "params*.yml"):
            for p in base.glob(pattern):
                if p.is_file():
                    key = str(p.resolve())
                    if key not in seen:
                        seen.add(key)
                        matches.append(p)
    if not matches:
        return None
    return sorted(matches, key=_version_sort_key, reverse=True)[0]


def resolve_params_path() -> Path:
    """Return the resolved params file path using the same logic as load_cfg()."""
    env_path = os.environ.get("TREND_TRACKER_PARAMS")
    if env_path:
        return Path(env_path)

    cfg_path = _resolve_params_path([ROOT, ROOT / "CONFIG"])
    if cfg_path is None:
        raise FileNotFoundError(
            "No params*.yaml or params*.yml file found in project root or CONFIG/"
        )
    return cfg_path


def load_cfg(path: str | Path | None = None) -> dict[str, Any]:
    """Load params.yaml (or any versioned params*.yaml) and return it as a nested dict.

    Resolution order when path is omitted:
        1. TREND_TRACKER_PARAMS environment variable (if set)
        2. Highest-versioned params*.yaml / params*.yml in {ROOT}
        3. Highest-versioned params*.yaml / params*.yml in {ROOT}/CONFIG

    Set TREND_TRACKER_PARAMS to an absolute path before launching Jupyter to
    use an alternate params file without modifying notebooks:
        export TREND_TRACKER_PARAMS=/path/to/custom_params.yaml

    Args:
        path: Explicit path to a YAML file. When provided, skips the
              resolution order entirely (including the env variable).

    Returns:
        Parsed YAML content as a dict.
    """
    if path is not None:
        cfg_path = Path(path)
    else:
        env_path = os.environ.get("TREND_TRACKER_PARAMS")
        if env_path:
            cfg_path = Path(env_path)
        else:
            cfg_path = _resolve_params_path([ROOT, ROOT / "CONFIG"])
            if cfg_path is None:
                raise FileNotFoundError(
                    "No params*.yaml or params*.yml file found in project root or CONFIG/"
                )

    with open(cfg_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def write_json(path: Path, payload: Any) -> Path:
    """Write payload to a JSON file, creating parent directories as needed.

    Uses deterministic serialisation (sorted keys, ensure_ascii=False,
    default=str for non-serialisable types) so manifests produce stable diffs.

    Args:
        path:    Destination file path.
        payload: JSON-serialisable object.

    Returns:
        The resolved Path that was written.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False, default=str)
    return path


def compute_md5(path: Path, chunk_size: int = 1024 * 1024) -> str | None:
    """Return the MD5 hex digest of a file, or None if the file does not exist.

    Args:
        path:       File to hash.
        chunk_size: Read chunk size in bytes.

    Returns:
        Lowercase hex digest string, or None.
    """
    path = Path(path)
    if not path.exists() or not path.is_file():
        return None
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(chunk_size), b""):
            h.update(chunk)
    return h.hexdigest()


def artifact_meta(path: Path, label: str | None = None) -> dict[str, Any]:
    """Build a provenance metadata dict for one artifact path.

    Args:
        path:  Path to the artifact (may or may not exist yet).
        label: Human-readable label; defaults to path.name.

    Returns:
        Dict with keys: label, path, exists, md5, size_bytes.
    """
    path = Path(path)
    exists = path.exists()
    return {
        "label": label or path.name,
        "path": str(path),
        "exists": exists,
        "md5": compute_md5(path) if exists and path.is_file() else None,
        "size_bytes": path.stat().st_size if exists else None,
    }


def build_output_path(
    subdir: str,
    fname: str,
    groupby_field: str | None = None,
    run_date: str | None = None,
    root: Path = ROOT,
) -> Path:
    """Resolve a canonical OUTPUTS path and create parent directories.

    Without groupby_field / run_date (NB01, NB02):
        OUTPUTS/{subdir}/{fname}
    With both (legacy grouped outputs):
        OUTPUTS/{groupby_field}/{run_date}/{subdir}/{fname}

    Args:
        subdir:        Subdirectory under OUTPUTS (e.g. "prepared", "enrichment").
        fname:         File name including extension.
        groupby_field: Optional grouping column name for run-scoped paths.
        run_date:      Optional YYYY-MM-DD string for run-scoped paths.
        root:          Project root; defaults to the module ROOT constant.

    Returns:
        Resolved Path with parent directories created.
    """
    root = Path(root)
    if groupby_field and run_date:
        p = root / "OUTPUTS" / groupby_field / run_date / subdir / fname
    else:
        p = root / "OUTPUTS" / subdir / fname
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


def build_run_output_path(
    subdir: str,
    fname: str,
    groupby_field: str,
    run_date: str,
    run_id: str,
    root: Path = ROOT,
) -> Path:
    """Resolve a run-scoped OUTPUTS path for NB03.

    Path shape: OUTPUTS/runs/{groupby_field}/{run_date}/{run_id}/{subdir}/{fname}

    Args:
        subdir:        Subdirectory (e.g. "analysis", "insights").
        fname:         File name including extension.
        groupby_field: Grouping column name.
        run_date:      YYYY-MM-DD string.
        run_id:        Unique run identifier from get_run_id().
        root:          Project root; defaults to ROOT.

    Returns:
        Resolved Path with parent directories created.
    """
    root = Path(root)
    p = (
        root
        / "OUTPUTS"
        / "runs"
        / str(groupby_field)
        / str(run_date)
        / str(run_id)
        / subdir
        / fname
    )
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


def outpath(
    subdir: str,
    fname: str,
    root: Path = ROOT,
    groupby_field: str | None = None,
    run_date: str | None = None,
) -> Path:
    """Backward-compatible alias for build_output_path().

    Deprecated: use build_output_path() or build_run_output_path() directly.
    This alias will be removed in a future release.
    """
    return build_output_path(
        subdir, fname, groupby_field=groupby_field, run_date=run_date, root=root
    )


def get_run_date() -> str:
    """Return today's date as a YYYY-MM-DD string for output path nesting."""
    return datetime.now().strftime("%Y-%m-%d")


# ── 3. Run identity & filter helpers ─────────────────────────────────────────


def canonicalize_filter_spec(
    filter_logic: str,
    filters: list[dict[str, Any]] | None,
) -> dict[str, Any]:
    """Return a canonical, hashable filter-spec object used for run IDs and logging.

    Keys within each filter rule are sorted so that equivalent specs always
    produce the same JSON string regardless of dict insertion order.

    Args:
        filter_logic: Combination logic; only "and" is currently supported.
        filters:      List of filter rule dicts. None is treated as [].

    Returns:
        Dict with keys: schema_version, filter_logic, filters.
    """
    filters = filters or []
    canonical = []
    for item in filters:
        if not isinstance(item, dict):
            raise ValueError(f"Filter entries must be dicts, got {type(item)}")
        canonical.append({k: item[k] for k in sorted(item.keys())})
    return {"schema_version": "v1", "filter_logic": filter_logic, "filters": canonical}


def get_filter_fields_key(filters: list[dict[str, Any]] | None) -> str:
    """Return a stable string key summarising which fields are filtered on.

    Used to keep run outputs from different filter scopes from colliding.

    Args:
        filters: List of filter rule dicts. None or [] returns "none".

    Returns:
        Double-underscore-joined sorted field names, or "none".
    """
    filters = filters or []
    fields = sorted(
        {str(f.get("field", "")).strip() for f in filters if str(f.get("field", "")).strip()}
    )
    return "none" if not fields else "__".join(fields)


def get_run_id(
    groupby_field: str,
    filter_spec: dict[str, Any] | None = None,
) -> str:
    """Build a human-readable, sortable run ID with an 8-char scope hash.

    Format: {YYYYMMDD_HHMMSS}_{groupby_field}_{hash}

    Args:
        groupby_field: The analysis grouping column.
        filter_spec:   Canonical filter spec from canonicalize_filter_spec().
                       None is treated as an empty filter spec.

    Returns:
        Run ID string.
    """
    scope = json.dumps(
        filter_spec or {"filters": [], "filter_logic": "and"},
        sort_keys=True,
        separators=(",", ":"),
    )
    scope_hash = hashlib.md5(scope.encode("utf-8")).hexdigest()[:8]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{timestamp}_{groupby_field}_{scope_hash}"


def validate_filter_spec(
    df: pd.DataFrame,
    filter_logic: str,
    filters: list[dict[str, Any]] | None,
) -> None:
    """Validate analysis filter configuration against the loaded dataframe.

    Raises ValueError on the first problem found. Intended to be called
    at notebook startup so errors surface early, before any LLM calls.

    Args:
        df:           The dataframe the filters will be applied to.
        filter_logic: Must be exactly "and".
        filters:      List of filter rule dicts. None or [] passes validation.

    Raises:
        ValueError: On unsupported filter_logic, missing fields, bad ops,
                    or missing required rule keys.
    """
    if filter_logic != "and":
        raise ValueError("analysis.filter_logic must be exactly 'and'")
    for rule in filters or []:
        if "field" not in rule:
            raise ValueError(f"Filter rule missing 'field': {rule}")
        field = rule["field"]
        if field not in df.columns:
            raise ValueError(f"Filter field '{field}' not found in dataframe columns")
        op = rule.get("op")
        if op not in {"eq", "in", "range", "is_null", "not_null"}:
            raise ValueError(f"Unsupported filter op '{op}' in rule: {rule}")
        if op == "eq" and "value" not in rule:
            raise ValueError(f"eq filter missing 'value': {rule}")
        if op == "in":
            values = rule.get("values")
            if not isinstance(values, list) or not values:
                raise ValueError(f"in filter requires non-empty 'values': {rule}")
        if op == "range" and "min" not in rule and "max" not in rule:
            raise ValueError(f"range filter requires 'min' and/or 'max': {rule}")


def _coerce_bound_for_series(series: pd.Series, value: Any) -> Any:
    """Coerce a range bound to match the dtype of the target series.

    Only datetime columns are auto-coerced; all other types are returned
    unchanged to avoid silent type mismatches.

    Args:
        series: The dataframe column the bound will be compared against.
        value:  The raw bound value from the filter rule.

    Returns:
        Coerced value (datetime for datetime columns, original value otherwise).
    """
    if pd.api.types.is_datetime64_any_dtype(series):
        return pd.to_datetime(value)
    return value


def apply_filters(
    df: pd.DataFrame,
    filter_logic: str,
    filters: list[dict[str, Any]] | None,
) -> tuple[pd.DataFrame, dict[str, Any]]:
    """Apply validated analysis filters and return the filtered dataframe plus a summary.

    Args:
        df:           Input dataframe.
        filter_logic: Combination logic; only "and" is supported.
        filters:      List of filter rule dicts. None or [] returns df unchanged.

    Returns:
        (filtered_df, summary_dict) where summary_dict contains row counts,
        retained percentage, fields checked, and filter_fields_key.
    """
    filters = filters or []
    validate_filter_spec(df, filter_logic, filters)

    if not filters:
        return df.copy(), {
            "filter_logic": filter_logic,
            "filters": [],
            "n_rules": 0,
            "input_row_count": int(len(df)),
            "output_row_count": int(len(df)),
            "dropped_row_count": 0,
            "retained_pct": 100.0 if len(df) else 0.0,
            "fields_checked": [],
            "no_rows_after_filter": False,
            "filter_fields_key": "none",
        }

    mask = pd.Series(True, index=df.index)
    for rule in filters:
        field, op, series = rule["field"], rule["op"], df[rule["field"]]
        if op == "eq":
            rule_mask = series == rule["value"]
        elif op == "in":
            rule_mask = series.isin(rule["values"])
        elif op == "range":
            rule_mask = pd.Series(True, index=df.index)
            if "min" in rule:
                rule_mask &= series >= _coerce_bound_for_series(series, rule["min"])
            if "max" in rule:
                rule_mask &= series <= _coerce_bound_for_series(series, rule["max"])
        elif op == "is_null":
            rule_mask = series.isna()
        elif op == "not_null":
            rule_mask = series.notna()
        else:
            raise ValueError(f"Unsupported filter op '{op}'")
        mask &= rule_mask.fillna(False)

    out = df.loc[mask].copy()
    retained_pct = round(len(out) / len(df) * 100, 2) if len(df) else 0.0
    return out, {
        "filter_logic": filter_logic,
        "filters": filters,
        "n_rules": len(filters),
        "input_row_count": int(len(df)),
        "output_row_count": int(len(out)),
        "dropped_row_count": int(len(df) - len(out)),
        "retained_pct": retained_pct,
        "fields_checked": [rule["field"] for rule in filters],
        "no_rows_after_filter": out.empty,
        "filter_fields_key": get_filter_fields_key(filters),
    }


# ── 4. Stage & pipeline manifests ────────────────────────────────────────────


def start_stage_manifest(
    stage_name: str,
    notebook_file: str,
    config_path: str = "params.yaml",
    run_id: str | None = None,
    group_by_field: str | None = None,
    filter_fields_key: str | None = None,
) -> dict[str, Any]:
    """Create the common stage-manifest skeleton shared by all three notebooks.

    The manifest starts with status="running" and is finalised by a later call
    to finalize_stage_manifest().

    Args:
        stage_name:       Human-readable name for this pipeline stage.
        notebook_file:    Actual filename of the calling notebook (for provenance).
        config_path:      Path to params.yaml relative to ROOT.
        run_id:           Run ID from get_run_id(), or None for NB01/NB02.
        group_by_field:   Grouping column name, or None for NB01/NB02.
        filter_fields_key: Key from get_filter_fields_key(), or None.

    Returns:
        Manifest dict ready to be passed to finalize_stage_manifest().
    """
    config_meta = artifact_meta(ROOT / config_path, label="config")
    return {
        "schema_version": "v1",
        "pipeline_version": PIPELINE_VERSION,
        "run_id": run_id,
        "group_by_field": group_by_field,
        "filter_fields_key": filter_fields_key,
        "stage_name": stage_name,
        "status": "running",
        "started_at": datetime.now().isoformat(),
        "completed_at": None,
        "duration_seconds": None,
        "notebook_file": notebook_file,
        "config_path": config_path,
        "config_md5": config_meta["md5"],
        "input_artifacts": [],
        "output_artifacts": [],
        "row_counts": {},
        "key_params": {},
        "warnings_count": 0,
        "warnings_path": None,
    }


def finalize_stage_manifest(
    manifest: dict[str, Any],
    output_path: Path,
    status: str,
    input_artifacts: list[dict[str, Any]] | None = None,
    output_artifacts: list[dict[str, Any]] | None = None,
    row_counts: dict[str, Any] | None = None,
    key_params: dict[str, Any] | None = None,
    warnings_path: Path | None = None,
) -> dict[str, Any]:
    """Finalize a stage manifest, compute duration, and write it to disk.

    Args:
        manifest:          Dict returned by start_stage_manifest().
        output_path:       Where to write the final JSON manifest.
        status:            "success" or "failure".
        input_artifacts:   List of artifact_meta() dicts for inputs.
        output_artifacts:  List of artifact_meta() dicts for outputs.
        row_counts:        Dict of labelled row counts for the QA record.
        key_params:        Dict of parameter values recorded for the QA record.
        warnings_path:     Path to the JSONL warnings file for this stage.

    Returns:
        The completed manifest dict (also written to output_path).
    """
    completed_at = datetime.now()
    started_at = datetime.fromisoformat(manifest["started_at"])
    manifest["status"] = status
    manifest["completed_at"] = completed_at.isoformat()
    manifest["duration_seconds"] = round((completed_at - started_at).total_seconds(), 2)
    manifest["input_artifacts"] = input_artifacts or []
    manifest["output_artifacts"] = output_artifacts or []
    manifest["row_counts"] = row_counts or {}
    manifest["key_params"] = key_params or {}
    if warnings_path is not None:
        manifest["warnings_path"] = str(warnings_path)
        manifest["warnings_count"] = get_warning_count(warnings_path)
    write_json(Path(output_path), manifest)
    return manifest


def build_pipeline_manifest(
    output_path: Path,
    run_id: str,
    run_date: str,
    group_by_field: str,
    filter_spec_path: Path,
    filter_summary_path: Path,
    stage_manifest_paths: list[Path],
    warnings_01_path: Path,
    warnings_02_path: Path,
    warnings_03_path: Path,
    final_outputs: dict[str, str],
    config_path: str = "params.yaml",
    filter_fields_key: str | None = None,
    status: str = "success",
) -> dict[str, Any]:
    """Persist the end-of-NB03 pipeline manifest covering all three stages.

    Args:
        output_path:           Where to write the pipeline manifest JSON.
        run_id:                Run ID from get_run_id().
        run_date:              YYYY-MM-DD string.
        group_by_field:        Grouping column name.
        filter_spec_path:      Path to the saved filter spec JSON.
        filter_summary_path:   Path to the saved filter summary JSON.
        stage_manifest_paths:  Paths to all three stage manifests.
        warnings_01_path:      NB01 warnings JSONL path.
        warnings_02_path:      NB02 warnings JSONL path.
        warnings_03_path:      NB03 warnings JSONL path.
        final_outputs:         Dict of labelled final output paths.
        config_path:           Path to params.yaml relative to ROOT.
        filter_fields_key:     Key from get_filter_fields_key().
        status:                "success" or "failure".

    Returns:
        The manifest dict (also written to output_path).
    """
    config_meta = artifact_meta(ROOT / config_path, label="config")
    payload = {
        "schema_version": "v1",
        "pipeline_version": PIPELINE_VERSION,
        "run_id": run_id,
        "group_by_field": group_by_field,
        "filter_fields_key": filter_fields_key,
        "run_date": run_date,
        "status": status,
        "created_at": datetime.now().isoformat(),
        "config_path": config_path,
        "config_md5": config_meta["md5"],
        "filter_spec_path": str(filter_spec_path),
        "filter_summary_path": str(filter_summary_path),
        "stage_manifests": [str(Path(p)) for p in stage_manifest_paths],
        # Individual warning paths kept for backward compatibility with consumers
        # that reference them by key. warnings_files is the canonical list form.
        "warnings_01_path": str(warnings_01_path),
        "warnings_02_path": str(warnings_02_path),
        "warnings_03_path": str(warnings_03_path),
        "warnings_files": [
            str(warnings_01_path),
            str(warnings_02_path),
            str(warnings_03_path),
        ],
        "final_outputs": final_outputs,
    }
    write_json(output_path, payload)
    return payload


# ── 5. Warning file helpers ───────────────────────────────────────────────────


def ensure_warning_file(path: Path) -> Path:
    """Create an empty JSONL warnings file if it does not already exist.

    Args:
        path: Target JSONL file path.

    Returns:
        The resolved Path.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        path.write_text("", encoding="utf-8")
    return path


def append_warning(
    path: Path,
    stage_name: str,
    code: str,
    message: str,
    severity: str = "warning",
    context: dict[str, Any] | None = None,
) -> None:
    """Append one structured warning record to a JSONL warnings file.

    Each record contains a UTC timestamp, stage name, severity, code,
    human-readable message, and an optional context dict for debugging.

    Args:
        path:       JSONL file path (created if missing).
        stage_name: Name of the pipeline stage emitting the warning.
        code:       Machine-readable warning code (e.g. "NMF_GROUP_SKIPPED").
        message:    Human-readable description.
        severity:   "warning" or "error".
        context:    Optional dict of additional debug fields.
    """
    ensure_warning_file(path)
    record = {
        "timestamp": datetime.now().isoformat(),
        "stage_name": stage_name,
        "severity": severity,
        "code": code,
        "message": message,
        "context": context or {},
    }
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False, default=str) + "\n")


def get_warning_count(path: Path) -> int:
    """Count non-empty lines in a JSONL warnings file.

    Args:
        path: JSONL file path.

    Returns:
        Number of warning records, or 0 if the file does not exist.
    """
    path = Path(path)
    if not path.exists():
        return 0
    with open(path, "r", encoding="utf-8") as f:
        return sum(1 for line in f if line.strip())


# ── 6. LLM client ────────────────────────────────────────────────────────────


def get_llm_client() -> OpenAI:
    """Build and return an OpenAI client with SSL verification disabled.

    SSL verification is disabled to work with the DonorsChoose proxy
    environment. This is a fixed infrastructure requirement, not a
    configurable option.

    Returns:
        Authenticated OpenAI client.

    Raises:
        ValueError: If OPENAI_API_KEY is not set in the environment.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("Missing OPENAI_API_KEY environment variable.")
    return OpenAI(api_key=api_key, http_client=httpx.Client(verify=False))


# ── 7. Ingest ─────────────────────────────────────────────────────────────────


def ingest(path: str | Path) -> pd.DataFrame:
    """Load a SQL extract CSV and parse the LISTAGG token string into a list.

    The SQL extract produces tokens as a comma-separated string (the result
    of a LISTAGG aggregation). This function splits that string into a proper
    Python list and parses known date columns.

    Args:
        path: Path to the CSV file.

    Returns:
        DataFrame with a 'tokens' column containing lists of strings.
    """
    raw = pd.read_csv(path)
    date_cols = [c for c in ["posted_date", "funded_date"] if c in raw.columns]
    df = pd.read_csv(path, parse_dates=date_cols) if date_cols else raw
    df["tokens"] = (
        df["tokens"]
        .fillna("")
        .str.split(",")
        .apply(lambda ts: [t.strip() for t in ts if t.strip()])
    )
    return df


# ── 8. Token helpers ─────────────────────────────────────────────────────────


def tokens_to_str(token_list: list[str] | None) -> str:
    """Join a token list into a space-separated string for sklearn vectorizers.

    Args:
        token_list: List of token strings, or None.

    Returns:
        Space-joined string, or "" for None / empty input.
    """
    return " ".join(token_list) if token_list else ""


def flat_freq(df: pd.DataFrame, col: str = "tokens") -> pd.Series:
    """Compute corpus-wide token frequency across all projects.

    Args:
        df:  DataFrame with a column containing token lists.
        col: Name of the token-list column.

    Returns:
        Series of (token → count) sorted descending by count.
    """
    return pd.Series([t for ts in df[col] for t in ts]).value_counts()


def token_doc_freq(df: pd.DataFrame) -> pd.Series:
    """Compute the number of distinct projects containing each token.

    Uses explode rather than a loop, so it scales to large corpora.

    Args:
        df: DataFrame with 'project_id' and 'tokens' columns.

    Returns:
        Series of (token → distinct project count) named "doc_count".
    """
    return (
        df[["project_id", "tokens"]]
        .explode("tokens")
        .rename(columns={"tokens": "token"})
        .drop_duplicates()
        .groupby("token")["project_id"]
        .nunique()
        .rename("doc_count")
    )


def normalize_tokens(tokens: list[str], lang: str = "en") -> list[str]:
    """Apply morphological normalization to a token list via simplemma.

    Handles inflected forms correctly (drives→drive, organized→organize).
    Domain terms and proper nouns that simplemma does not recognise are
    returned unchanged; the consolidation map handles residual cases.

    Args:
        tokens: List of raw token strings.
        lang:   ISO 639-1 language code for simplemma. Defaults to "en".

    Returns:
        List of normalized token strings (same length as input).
    """
    return [simplemma.lemmatize(t, lang=lang) for t in tokens]


# ── 9. Consolidation helpers ─────────────────────────────────────────────────


def _auto_replacement(
    original: str,
    token_set: set[str],
    known_fixes: dict[str, str],
) -> str:
    """Return a replacement string for a token, or '' to leave it unchanged.

    Checks in order:
        1. known_fixes dict (simplemma truncation artifacts and other corrections).
        2. Singular-s stripping: map 'words' → 'word' when 'word' is in vocab.
        3. Singular-es stripping: map 'wishes' → 'wish' when 'wish' is in vocab.

    Args:
        original:    The token to check.
        token_set:   Set of all tokens in the top-N vocabulary window.
        known_fixes: Dict of {bad_stem: correct_form} loaded from
                     CONFIG/known_lemma_fixes.yaml.

    Returns:
        Replacement string, or '' if no replacement applies.
    """
    if original in known_fixes:
        return known_fixes[original]
    if original.endswith("s") and not original.endswith("ss") and len(original) > 3:
        singular = original[:-1]
        if singular in token_set:
            return singular
    if original.endswith("es") and len(original) > 4:
        singular = original[:-2]
        if singular in token_set:
            return singular
    return ""


def build_consolidation_candidates(
    df: pd.DataFrame,
    top_n: int = 1000,
    known_fixes: dict[str, str] | None = None,
) -> pd.DataFrame:
    """Build a consolidation candidate table from the top-N vocabulary tokens.

    For each of the top_n tokens by corpus frequency, auto-assigns a
    replacement (via _auto_replacement) and flags likely truncation artifacts
    (via _TRUNC_RE). The resulting CSV is the human-reviewable consolidation map
    used in NB01 Step 3.

    The known_fixes dict should be loaded from CONFIG/known_lemma_fixes.yaml
    (see params.yaml: preprocess.known_lemma_fixes_path). If not supplied,
    only plural-pair detection runs; KNOWN_FIXES corrections are skipped.

    Args:
        df:          DataFrame with a 'tokens' column (token lists).
        top_n:       Number of top-frequency tokens to include. Default 1000.
        known_fixes: Dict of {bad_stem: correct_form}. None treated as {}.

    Returns:
        DataFrame with columns: original, freq, flag_type, replacement, notes.
    """
    if known_fixes is None:
        known_fixes = {}

    vocab = flat_freq(df).head(top_n).reset_index()
    vocab.columns = ["original", "freq"]
    vocab["flag_type"] = vocab["original"].apply(
        lambda w: "truncation" if len(w) <= 8 and _TRUNC_RE.search(w) else ""
    )
    token_set = set(vocab["original"])
    vocab["replacement"] = vocab["original"].apply(
        lambda w: _auto_replacement(w, token_set, known_fixes)
    )
    vocab["notes"] = "auto"
    return vocab


# ── 10. Analysis helpers ──────────────────────────────────────────────────────


def make_vec(
    min_df: int | float,
    max_df: int | float,
    ngram_range: tuple[int, int],
) -> TfidfVectorizer:
    """Instantiate a TF-IDF vectorizer with the pipeline's token pattern.

    The token pattern allows underscore and hyphen within tokens to preserve
    injected enrichment tokens (e.g. __framing_urgency__) and hyphenated terms.

    Args:
        min_df:       Minimum document frequency (int count or float fraction).
        max_df:       Maximum document frequency (int count or float fraction).
        ngram_range:  (min_n, max_n) tuple.

    Returns:
        Configured but unfitted TfidfVectorizer.
    """
    return TfidfVectorizer(
        min_df=min_df,
        max_df=max_df,
        ngram_range=ngram_range,
        token_pattern=r"(?u)\b[a-z][a-z_\-]*\b",
    )


def add_bin(df: pd.DataFrame, bins: list[dict[str, Any]]) -> pd.DataFrame:
    """Label each project with the first matching analyst-defined date bin.

    Bins are matched in order against posted_date. Projects that match no
    bin receive None. Bins are defined in params.yaml under analysis.bins.

    Args:
        df:   DataFrame with a 'posted_date' column.
        bins: List of {name, start, end} dicts.

    Returns:
        Copy of df with an added 'bin' column.
    """
    df = df.copy()
    df["bin"] = None
    for b in bins:
        mask = (df["posted_date"] >= b["start"]) & (df["posted_date"] <= b["end"])
        df.loc[mask & df["bin"].isna(), "bin"] = b["name"]
    return df


def group_key(keys: Any, group_cols: list[str]) -> dict[str, Any]:
    """Normalise a groupby key (scalar or tuple) to a dict.

    pandas returns a scalar key for single-column groupby and a tuple for
    multi-column groupby. This function normalises both to a dict so
    downstream code does not need to branch.

    Args:
        keys:       The key returned by DataFrame.groupby().
        group_cols: List of column names used in the groupby.

    Returns:
        Dict mapping column name → key value.
    """
    return dict(zip(group_cols, keys if isinstance(keys, tuple) else [keys]))


def build_project_topic_bridge(
    weights_df: pd.DataFrame,
    groupby_field: str,
    threshold: float,
) -> pd.DataFrame:
    """Build the project-topic bridge table from NMF weight outputs.

    For each project, computes topic_share = weight / sum(weights across topics
    for that project), then retains only rows where topic_share >= threshold.

    Args:
        weights_df:    DataFrame with columns: {groupby_field}, topic_id,
                       project_id, weight.
        groupby_field: The analysis grouping column.
        threshold:     Minimum topic_share for a project-topic link to be kept.
                       Sourced from analysis.topic_assignment_threshold.

    Returns:
        DataFrame with columns: topic_key, project_id, {groupby_field},
        topic_id, weight, topic_share.

    Raises:
        ValueError: If required columns are missing from weights_df.
    """
    required = {groupby_field, "topic_id", "project_id", "weight"}
    missing = required - set(weights_df.columns)
    if missing:
        raise ValueError(f"weights_df missing required columns: {sorted(missing)}")

    totals = (
        weights_df.groupby([groupby_field, "project_id"])["weight"]
        .sum()
        .rename("total_weight")
        .reset_index()
    )
    merged = weights_df.merge(
        totals, on=[groupby_field, "project_id"], how="left", validate="many_to_one"
    )
    merged["topic_share"] = 0.0
    nonzero = merged["total_weight"].fillna(0) > 0
    merged.loc[nonzero, "topic_share"] = (
        merged.loc[nonzero, "weight"] / merged.loc[nonzero, "total_weight"]
    )
    linked = merged[merged["topic_share"] >= threshold].copy()
    linked["topic_key"] = (
        groupby_field
        + "="
        + linked[groupby_field].astype(str)
        + "|topic="
        + linked["topic_id"].astype(str)
    )
    return linked[
        ["topic_key", "project_id", groupby_field, "topic_id", "weight", "topic_share"]
    ]


def slugify_group_value(value: str, max_len: int = 64) -> str:
    """Convert a group value to a safe filesystem component for output paths.

    Strips accents, replaces non-word characters with underscores, collapses
    consecutive underscores, and truncates to max_len.

    Args:
        value:   Arbitrary group value string.
        max_len: Maximum length of the returned slug. Default 64.

    Returns:
        Safe lowercase slug, or "unknown" for empty/unrepresentable input.
    """
    value = str(value)
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    value = re.sub(r"[^\w]", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value[:max_len] if value else "unknown"


def load_essay_snippet_lookup(
    project_ids: list[Any],
    data_dir: Path | None = None,
    max_chars: int = 300,
) -> dict[Any, str]:
    """Lazily load essay text snippets for a set of project IDs.

    Scans DATA/project_essay*.csv files in order, returning early once all
    requested IDs are found. Only loads the project_id and text columns to
    keep memory usage low.

    Args:
        project_ids: List of project IDs to look up.
        data_dir:    Directory containing project_essay*.csv files.
                     Defaults to {ROOT}/DATA.
        max_chars:   Maximum character length of each returned snippet.

    Returns:
        Dict mapping project_id → truncated essay text for found IDs.
        IDs not found in any file are absent from the dict.
    """
    data_dir = Path(data_dir) if data_dir is not None else ROOT / "DATA"
    needed = set(project_ids)
    if not needed:
        return {}

    lookup: dict[Any, str] = {}
    essay_files = sorted(data_dir.glob("project_essay*.csv"))
    text_col_candidates = ["essay", "essay_text", "full_text", "project_essay", "text"]

    for fpath in essay_files:
        try:
            cols = pd.read_csv(fpath, nrows=0).columns.tolist()
        except Exception:
            continue
        text_col = next((c for c in text_col_candidates if c in cols), None)
        if not text_col or "project_id" not in cols:
            continue
        for chunk in pd.read_csv(fpath, usecols=["project_id", text_col], chunksize=200_000):
            sub = chunk[chunk["project_id"].isin(needed - set(lookup.keys()))]
            if sub.empty:
                continue
            for _, row in sub.iterrows():
                text = re.sub(r"\s+", " ", str(row.get(text_col, "") or "")).strip()
                if text:
                    lookup[row["project_id"]] = text[:max_chars]
            if len(lookup) == len(needed):
                return lookup
    return lookup


# ── 11. Quality ───────────────────────────────────────────────────────────────

# Module-level stopword set used as the default fallback in quality_report().
# These are terms that should never survive preprocessing in a clean corpus.
#
# NOTE: This constant is the live fallback until quality_report() is updated
# to load from params.yaml → quality.stopword_violation_list. Keep this set
# in sync with that YAML list. Once loading from params is wired in, this
# constant can be removed.
HARD_STOPWORDS: set[str] = {
    "the", "and", "for", "with", "are", "was", "were", "been", "this", "that",
    "from", "they", "them", "their", "will", "would", "should", "could", "have",
    "has", "had", "use", "using", "used", "make", "makes", "making", "get",
    "gets", "getting", "help", "helps", "project", "students", "student",
    "classroom", "learning", "school", "teacher", "teachers", "education",
    "grade", "grades", "materials", "supplies", "tools", "resources",
    "funded", "funding", "donors", "donor",
}


def quality_report(
    df: pd.DataFrame,
    label: str,
    doc_freq: pd.Series | None = None,
    matrices: dict[str, Any] | None = None,
    save_path: Path | None = None,
    stopwords: list[str] | set[str] | None = None,
) -> dict[str, Any]:
    """Generate and print a corpus quality snapshot.

    Checks token distribution, vocabulary size, and stopword violations.
    Used at pipeline checkpoints (NB01 Step 6, NB03 Step 2) to gate progress
    before expensive LLM calls.

    Args:
        df:        DataFrame with a 'tokens' column.
        label:     Checkpoint label for display and JSON output (e.g. "cp1").
        doc_freq:  Optional Series of document frequencies for additional stats.
        matrices:  Optional dict of {name: scipy sparse matrix} to report on.
        save_path: Optional path to write the quality stats JSON.
        stopwords: Set or list of stopwords to check against. When None,
                   falls back to the HARD_STOPWORDS module constant. Intended
                   to be loaded from params.yaml → quality.stopword_violation_list
                   once that wiring is in place.

    Returns:
        Dict of quality statistics including stopword gate result.
    """
    active_stopwords = set(stopwords) if stopwords is not None else HARD_STOPWORDS
    freq = flat_freq(df)
    stops = [t for t in freq.head(200).index if t in active_stopwords]

    stats: dict[str, Any] = {
        "checkpoint": label,
        "timestamp": datetime.now().isoformat(),
        "n_projects": len(df),
        "token_count_distribution": df["tokens"].apply(len).describe().to_dict(),
        "vocab": {
            "unique": int(len(freq)),
            "total": int(freq.sum()),
            "stopword_violations": stops,
        },
        "gates": {"no_stopwords": not stops, "violations": stops},
    }

    if doc_freq is not None and len(doc_freq):
        stats["doc_freq"] = {
            "retained": int(len(doc_freq)),
            "min": int(doc_freq.min()),
            "max": int(doc_freq.max()),
            "median": float(doc_freq.median()),
        }

    if matrices:
        stats["matrices"] = {
            k: {
                "shape": list(X.shape),
                "nnz": int(X.nnz),
                "sparsity": round(1 - X.nnz / (X.shape[0] * X.shape[1] or 1), 4),
            }
            for k, X in matrices.items()
        }

    tc = stats["token_count_distribution"]
    print(f"\n{'=' * 55}  [{label}]")
    print(f"  Projects : {stats['n_projects']:,}")
    print(f"  Tok/proj : min={tc['min']:.0f}  p50={tc['50%']:.0f}  max={tc['max']:.0f}")
    print(f"  Vocab    : {stats['vocab']['unique']:,} unique tokens")
    if matrices:
        for k, m in stats["matrices"].items():
            print(f"  {k:20s}: shape={m['shape']}  sparsity={m['sparsity']:.3f}")
    print(f"  Stopwords: {'PASS' if not stops else 'FAIL — ' + str(stops)}")
    print(f"{'=' * 55}\n")

    if save_path:
        write_json(save_path, stats)
    return stats


# ── 12. TF-IDF & NMF helpers ─────────────────────────────────────────────────


def cat_tfidf_slice(
    idx: Any,
    df_index: Any,
    X_full: Any,
    feat: Any,
    idf_vals: Any,
    top_n: int,
) -> pd.DataFrame:
    """Score one group slice against the rest of the corpus using a shared TF-IDF matrix.

    The vectorizer is fit once on the full corpus; this function only
    indexes into the resulting matrix, so there is no per-group refitting.

    Args:
        idx:       Index labels for the focal group's rows.
        df_index:  Full dataframe index (used to compute the complement set).
        X_full:    Full corpus TF-IDF sparse matrix.
        feat:      Feature name array from vectorizer.get_feature_names_out().
        idf_vals:  IDF value array from vectorizer.idf_.
        top_n:     Number of highest-TF-IDF terms to return.

    Returns:
        DataFrame with columns: token, tf, idf, tfidf, prevalence,
        contrast, project_count — sorted by tfidf descending.
    """
    rest_idx = df_index.difference(idx)
    X_cat = X_full[idx.tolist()]
    X_rest = X_full[rest_idx.tolist()]

    cat_prev = (X_cat > 0).mean(axis=0).A1
    rest_prev = (X_rest > 0).mean(axis=0).A1 if len(rest_idx) else np.zeros(len(feat))
    tf = X_cat.mean(axis=0).A1

    return pd.DataFrame({
        "token": feat,
        "tf": tf,
        "idf": idf_vals,
        "tfidf": tf * idf_vals,
        "prevalence": cat_prev,
        "contrast": cat_prev - rest_prev,
        "project_count": (X_cat > 0).sum(axis=0).A1.astype(int),
    }).nlargest(top_n, "tfidf")


def choose_n_components(
    n_docs: int,
    retained_vocab: int,
    base_n_components: int,
    slice_rules: dict[str, Any],
) -> int:
    """Choose an NMF topic count using corpus-size and vocabulary-size heuristics.

    Three caps are applied and the minimum is taken:
        doc_cap   — prevents too many topics relative to group size.
                    Uses slice_rules["min_projects_per_topic"] (default 15)
                    as the divisor: at least that many projects per topic.
        vocab_cap — prevents topics outnumbering the retained vocabulary.
        topic_cap — respects base_n_components or small_slice_topic_cap
                    when small-slice mode is active.

    Small-slice mode is activated by the notebook when the median group size
    falls below slice_rules["small_slice_cutoff"]; the notebook injects
    slice_rules["small_slice_mode"] = True before calling nmf_one().

    Args:
        n_docs:            Number of documents in this slice.
        retained_vocab:    Number of features retained by the TF-IDF vectorizer.
        base_n_components: Base topic count from params.yaml → nmf.n_components.
        slice_rules:       Dict from params.yaml → analysis.slice_rules.

    Returns:
        Final NMF n_components value (minimum 4).
    """
    min_ppt = slice_rules.get("min_projects_per_topic", 15)
    doc_cap = max(4, n_docs // min_ppt)
    vocab_cap = max(4, retained_vocab // 8)
    topic_cap = (
        slice_rules["small_slice_topic_cap"]
        if slice_rules.get("small_slice_mode", False)
        else base_n_components
    )
    return max(4, min(base_n_components, doc_cap, vocab_cap, topic_cap))


def nmf_one(
    docs: list[str],
    ct_cfg: dict[str, Any],
    cn_cfg: dict[str, Any],
    base_n_components: int,
    slice_rules: dict[str, Any],
) -> tuple[pd.DataFrame | None, Any | None, dict[str, Any]]:
    """Fit one NMF slice for a single group and return topics, weights, and metadata.

    Returns early (with None, None, skip_meta) when the slice does not meet
    minimum vocabulary or matrix density requirements. Skip reasons are
    recorded in the returned metadata dict for the warnings log.

    Args:
        docs:              List of token strings (one per project).
        ct_cfg:            TF-IDF config dict (params.yaml → tfidf).
        cn_cfg:            NMF config dict (params.yaml → nmf).
        base_n_components: Base topic count (params.yaml → nmf.n_components).
        slice_rules:       Slice eligibility rules (params.yaml → analysis.slice_rules).

    Returns:
        (topics_df, W, meta) where:
            topics_df  — DataFrame with topic_id, top_terms, top_weights columns,
                         or None on skip.
            W          — NMF weight matrix (n_docs × n_topics), or None on skip.
            meta       — Dict with n_components_used, retained_vocab,
                         nonzero_tfidf_nnz, and optionally skip_reason.
    """
    vec = make_vec(
        ct_cfg["min_df"],
        ct_cfg["max_df"],
        tuple(ct_cfg.get("ngram_range", [1, 1])),
    )
    X = vec.fit_transform(docs)
    retained_vocab = int(X.shape[1])
    nonzero_tfidf_nnz = int(X.nnz)

    if retained_vocab < slice_rules["min_retained_vocab"]:
        return None, None, {
            "skip_reason": "low_retained_vocab",
            "retained_vocab": retained_vocab,
            "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
        }
    if nonzero_tfidf_nnz < slice_rules["min_tfidf_nnz"]:
        return None, None, {
            "skip_reason": "low_tfidf_nnz",
            "retained_vocab": retained_vocab,
            "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
        }

    n_components_used = choose_n_components(
        n_docs=len(docs),
        retained_vocab=retained_vocab,
        base_n_components=base_n_components,
        slice_rules=slice_rules,
    )
    if retained_vocab < n_components_used:
        return None, None, {
            "skip_reason": "vocab_below_topic_count",
            "retained_vocab": retained_vocab,
            "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
            "n_components_used": n_components_used,
        }

    model = NMF(
        n_components=n_components_used,
        random_state=cn_cfg["random_seed"],
        init="nndsvd",
        max_iter=cn_cfg["max_iter"],
    )
    W = model.fit_transform(X)
    vocab = vec.get_feature_names_out()

    rows = []
    for i, comp in enumerate(model.components_):
        idx = comp.argsort()[::-1][: cn_cfg["top_words"]]
        rows.append({
            "topic_id": i,
            "top_terms": vocab[idx].tolist(),
            "top_weights": comp[idx].tolist(),
        })

    return pd.DataFrame(rows), W, {
        "n_components_used": n_components_used,
        "retained_vocab": retained_vocab,
        "nonzero_tfidf_nnz": nonzero_tfidf_nnz,
    }


# ── 13. Enrichment helpers ────────────────────────────────────────────────────


def gate_cluster(
    cid: int,
    terms: list[str],
    *,
    client: OpenAI,
    model: str,
    system_prompt: str,
    prompt_template: str,
    retries: int = 2,
) -> dict[str, Any]:
    """Run the LLM coherence gate for one enrichment cluster (NB02 Pass A).

    Calls the model once per cluster to classify it as inject / split / discard.
    On failure after all retries, returns a discard result with the error reason.

    The prompt_template must accept {cid} and {terms} format keys.
    Terms are capped at the first 20 before formatting.

    Args:
        cid:             Cluster ID (integer label from agglomerative clustering).
        terms:           List of vocabulary terms in this cluster.
        client:          OpenAI client from get_llm_client().
        model:           Model name (params.yaml → models.gate).
        system_prompt:   System prompt for the gating call.
        prompt_template: User prompt template with {cid} and {terms} placeholders.
        retries:         Maximum retry attempts on transient errors.

    Returns:
        Dict with keys: action, primary_category, subcategory, split_into, reasoning.
    """
    prompt = prompt_template.format(cid=cid, terms=terms[:20])
    for attempt in range(retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
            )
            return json.loads(resp.choices[0].message.content.strip())
        except Exception as e:
            if attempt < retries:
                _time.sleep(2 ** attempt)
            else:
                return {
                    "action": "discard",
                    "primary_category": None,
                    "subcategory": None,
                    "split_into": [],
                    "reasoning": f"API error after {retries} retries: {e}",
                }


def classify_batch(
    terms_batch: list[str],
    *,
    client: OpenAI,
    model: str,
    system_prompt: str,
    prompt_template: str,
    taxonomy_ref: str,
    valid_categories: set[str],
    retries: int = 2,
) -> dict[str, str | None]:
    """Classify a batch of vocabulary terms against the framing taxonomy (NB02 Pass B).

    Sends terms_batch to the model with the taxonomy reference embedded in the
    prompt. Filters returned categories against valid_categories so that model
    hallucinations never reach the injection map.

    The prompt_template must accept {taxonomy} and {terms} format keys.

    Args:
        terms_batch:      List of vocabulary terms to classify.
        client:           OpenAI client from get_llm_client().
        model:            Model name (params.yaml → models.classify).
        system_prompt:    System prompt for the classification call.
        prompt_template:  User prompt template with {taxonomy} and {terms} keys.
        taxonomy_ref:     Pre-formatted taxonomy reference string (built in NB02).
        valid_categories: Set of legal category names from the loaded taxonomy.
        retries:          Maximum retry attempts on transient errors.

    Returns:
        Dict mapping term → category string (or None for unclassified terms).
        Only entries with None or a valid category are included.
    """
    prompt = prompt_template.format(taxonomy=taxonomy_ref, terms=terms_batch)
    for attempt in range(retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
            )
            raw = json.loads(resp.choices[0].message.content.strip())
            return {k: v for k, v in raw.items() if v is None or v in valid_categories}
        except Exception as e:
            if attempt < retries:
                _time.sleep(2 ** attempt)
            else:
                print(f"    Batch error: {e}")
                return {t: None for t in terms_batch}


def inject_tokens(
    token_list: list[str],
    lookup: dict[str, list[str]],
) -> list[str]:
    """Append injection tokens to a project's token list (NB02 Pass C).

    For each token in token_list, looks up any enrichment tokens to inject
    via the lookup dict. Injected tokens are deduplicated (preserving first
    occurrence order) and appended after the original tokens.

    Original tokens are never modified — injection is strictly additive.

    Args:
        token_list: Original token list for one project.
        lookup:     Dict mapping source token → list of enrichment tokens
                    to inject (built from semantic_map and framing_map CSVs).

    Returns:
        Original token list with deduplicated enrichment tokens appended,
        or the unmodified original list if no lookup matches were found.
    """
    extra = []
    for t in token_list:
        if t in lookup:
            extra.extend(lookup[t])
    if not extra:
        return token_list
    return token_list + list(dict.fromkeys(extra))


# ── 14. Topic labeling helpers ────────────────────────────────────────────────


def _norm_group_value(value: Any) -> str:
    """Normalise a group value to a lowercase stripped string for comparisons."""
    return str(value or "").strip().casefold()


def _safe_topic_id(value: Any) -> int:
    """Coerce a topic_id to int, returning -1 on failure."""
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return -1


def build_input(
    t_row: Any,
    weights_df: pd.DataFrame,
    pid_text: pd.Series,
    groupby_field: str,
    n_representative: int,
    top_terms_in_prompt: int,
) -> dict[str, Any]:
    """Build one topic-labeling payload from a topic row and top representative projects.

    Selects the highest-weight projects for this topic to use as representative
    snippets, then formats unigrams, bigrams, and NMF terms separately to give
    the model the best signal for label generation.

    Args:
        t_row:               Row from topics_df for one topic.
        weights_df:          NMF weight bridge DataFrame.
        pid_text:            Series mapping project_id → essay snippet.
        groupby_field:       The analysis grouping column.
        n_representative:    Number of representative snippets to include.
        top_terms_in_prompt: Maximum number of terms to include per term type.

    Returns:
        Dict with keys: group_value, topic_id, bin_line, unigrams, bigrams,
        nmf_terms, snippets — ready to format into a prompt template.
    """
    terms = t_row["top_terms"]
    key_cols = [groupby_field] + (["bin"] if "bin" in t_row.index else [])
    mask = weights_df["topic_id"] == t_row["topic_id"]
    for col in key_cols:
        mask &= weights_df[col] == t_row[col]

    rep_pids = (
        weights_df[mask]
        .sort_values("weight", ascending=False)["project_id"]
        .tolist()[:n_representative]
    )

    n_uni = top_terms_in_prompt
    n_bi = max(2, top_terms_in_prompt // 2)
    n_nmf = top_terms_in_prompt
    return {
        "group_value": t_row[groupby_field],
        "topic_id": int(t_row["topic_id"]),
        "bin_line": (
            f"\nBin: {t_row['bin']}"
            if "bin" in t_row.index and pd.notna(t_row.get("bin"))
            else ""
        ),
        "unigrams": ", ".join([x for x in terms if " " not in x][:n_uni]),
        "bigrams": ", ".join([x for x in terms if " " in x][:n_bi]),
        "nmf_terms": ", ".join(terms[:n_nmf]),
        "snippets": "\n".join(f"- {pid_text.get(p, '')}" for p in rep_pids),
    }


def _make_label_error(
    inp: dict[str, Any],
    raw_text: str,
    code: str,
    model_labeling: str,
    groupby_field: str,
    error_text: str | None = None,
) -> dict[str, Any]:
    """Return a structured error object for topic label failures.

    Used by _label_with_retry() to produce a consistently-shaped error record
    that can be stored alongside successful label results and filtered later.

    Args:
        inp:           The labeling input dict from build_input().
        raw_text:      Raw model response text (may be empty or malformed).
        code:          Machine-readable error code.
        model_labeling: Model name used for the failed call.
        groupby_field: The analysis grouping column.
        error_text:    String representation of the exception, if any.

    Returns:
        Dict with parse_error=True and standard label-result shape.
    """
    return {
        "raw": raw_text,
        "parse_error": True,
        "error_code": code,
        "error": error_text,
        "model": model_labeling,
        "timestamp": datetime.now().isoformat(),
        groupby_field: inp["group_value"],
        "topic_id": inp["topic_id"],
    }


def _label_with_retry(
    inp: dict[str, Any],
    *,
    client: OpenAI,
    model_labeling: str,
    system_prompt: str,
    user_prompt_template: str,
    groupby_field: str,
    warnings_path: Path,
    max_retries: int = 3,
    stage_name: str = "03_insights_generation",
) -> dict[str, Any]:
    """Call the labeling model with retry logic and structured error handling.

    Retries on JSON parse failures (backoff) and on rate-limit / timeout errors.
    Any other exception is treated as fatal for this topic (no retry).

    Args:
        inp:                  Labeling input dict from build_input().
        client:               OpenAI client.
        model_labeling:       Model name (params.yaml → models.labeling).
        system_prompt:        System prompt for the labeling call.
        user_prompt_template: User prompt template; formatted with inp as kwargs.
        groupby_field:        The analysis grouping column.
        warnings_path:        JSONL file for recording failures.
        max_retries:          Maximum retry attempts. Sourced from
                              params.yaml → llm.max_retries once wired.
        stage_name:           Stage name recorded in warning entries.

    Returns:
        Parsed label dict on success, or a _make_label_error() dict on failure.
    """
    text = ""
    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model_labeling,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt_template.format(**inp)},
                ],
            )
            text = resp.choices[0].message.content.strip()
            obj = json.loads(text)
            obj[groupby_field] = str(inp["group_value"])
            obj["topic_id"] = int(inp["topic_id"])
            obj["model"] = model_labeling
            obj["timestamp"] = datetime.now().isoformat()
            return obj

        except json.JSONDecodeError as e:
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
                continue
            append_warning(
                warnings_path, stage_name, "LABELING_PARSE_FAILURE",
                f"JSON parse failure for {inp['group_value']} / topic {inp['topic_id']}",
                context={"group": inp["group_value"], "topic_id": inp["topic_id"], "error": str(e)},
            )
            return _make_label_error(inp, text, "LABELING_PARSE_FAILURE",
                                     model_labeling, groupby_field, str(e))

        except (_openai.RateLimitError, _openai.APITimeoutError) as e:
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
            else:
                append_warning(
                    warnings_path, stage_name, "LABELING_API_FAILURE",
                    f"API failure after retries for {inp['group_value']} / topic {inp['topic_id']}",
                    context={"group": inp["group_value"], "topic_id": inp["topic_id"], "error": str(e)},
                )
                return _make_label_error(inp, text or str(e), "LABELING_API_FAILURE",
                                         model_labeling, groupby_field, str(e))

        except Exception as e:
            append_warning(
                warnings_path, stage_name, "LABELING_API_FAILURE",
                f"Unexpected error for {inp['group_value']} / topic {inp['topic_id']}",
                context={"group": inp["group_value"], "topic_id": inp["topic_id"], "error": str(e)},
            )
            return _make_label_error(inp, text or str(e), "LABELING_API_FAILURE",
                                     model_labeling, groupby_field, str(e))


# ── 15. Synthesis helpers ─────────────────────────────────────────────────────


def clean_label(text: Any) -> str:
    """Remove injected enrichment token markers from a label string.

    Strips patterns like __framing_urgency__ and __cat_marine_biology__ so
    that injected token names never appear in synthesis prompts or report text.

    Args:
        text: Any value; coerced to str before processing.

    Returns:
        Cleaned string with all __token_name__ patterns removed.
    """
    return re.sub(r"__[a-z_]+__\s*", "", str(text)).strip()


def build_topic_lines(
    df: pd.DataFrame,
    groupby_field: str,
    group: Any | None = None,
    top_terms_count: int = 4,
) -> str:
    """Render labeled topics into the line-oriented prompt format used in NB03.

    One line per topic, format:
        {group} | topic {id} | label: {label} | coherence: {flag} | [top_terms: ...] | description: {desc}

    Args:
        df:              labels_df from the labeling step.
        groupby_field:   The analysis grouping column.
        group:           If provided, filter to this group value only.
        top_terms_count: Maximum number of top terms to include per topic line.

    Returns:
        Newline-joined string of topic lines ready for prompt insertion.
    """
    if group is not None:
        df = df[df[groupby_field] == group]

    def _fmt_terms(val: Any, n: int) -> str:
        """Parse top_terms (list, JSON string, or CSV string) and return top n."""
        if n <= 0:
            return ""
        if isinstance(val, list):
            terms = [str(x).strip() for x in val if str(x).strip()]
        elif pd.isna(val):
            terms = []
        else:
            s = str(val).strip()
            if not s:
                terms = []
            else:
                try:
                    parsed = json.loads(s)
                    terms = [str(x).strip() for x in parsed if str(x).strip()] if isinstance(parsed, list) else [x.strip() for x in s.split(",") if x.strip()]
                except Exception:
                    terms = [x.strip() for x in s.split(",") if x.strip()]
        terms = [clean_label(t) for t in terms[:n]]
        return ", ".join(t for t in terms if t)

    lines = []
    for _, row in df.iterrows():
        top_terms_str = _fmt_terms(row.get("top_terms"), top_terms_count)
        line = (
            f"  {row[groupby_field]} | topic {row.topic_id} | "
            f"label: {clean_label(row.proposed_label)} | "
            f"coherence: {row.coherence_flag} | "
        )
        if top_terms_str:
            line += f"top_terms: {top_terms_str} | "
        line += f"description: {clean_label(row.description)}"
        lines.append(line)
    return "\n".join(lines)


def build_per_group_prompt(
    group: Any,
    group_description: str,
    topic_lines_text: str,
    per_group_instructions: str,
) -> str:
    """Build the per-group synthesis prompt body.

    Args:
        group:                  Group value (e.g. category name).
        group_description:      Human-readable description from group_descriptions.
        topic_lines_text:       Output of build_topic_lines() for this group.
        per_group_instructions: Instruction block appended after the topic list.

    Returns:
        Formatted prompt string ready for _call_with_retry().
    """
    return f"""
Below is a list of NMF topics discovered from teacher project request essays on DonorsChoose
for a single group: {group} ({group_description}).
Each topic represents a cluster of essays with similar language, framing, and request patterns.
{topic_lines_text}

{per_group_instructions}
""".strip()


def _call_with_retry(
    prompt: str,
    *,
    client: OpenAI,
    model_name: str,
    system_prompt: str,
    max_retries: int = 3,
) -> str | None:
    """Generic LLM caller with exponential backoff retry.

    Used for synthesis and cross-group analysis calls where the response
    is plain text (not JSON), so parse errors are not a retry trigger.

    Args:
        prompt:       User prompt string.
        client:       OpenAI client.
        model_name:   Model name.
        system_prompt: System prompt string.
        max_retries:  Maximum retry attempts on rate-limit or timeout errors.
                      Sourced from params.yaml → llm.max_retries once wired.

    Returns:
        Model response text stripped of leading/trailing whitespace,
        or None if a non-retryable error occurs.
    """
    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
            )
            return resp.choices[0].message.content.strip()
        except (_openai.RateLimitError, _openai.APITimeoutError):
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
            else:
                raise
        except Exception as e:
            print(f"Non-retryable error: {e}")
            return None


def synthesize_one_group(
    group: Any,
    *,
    labels_df: pd.DataFrame,
    groupby_field: str,
    group_description: str,
    per_group_instructions: str,
    client: OpenAI,
    model_name: str,
    system_prompt: str,
    warnings_path: Path,
    outpath_func: Callable[[str, str], Path],
    synthesis_top_terms_count: int = 4,
    max_retries: int = 3,
    stage_name: str = "03_insights_generation",
) -> tuple[Any, str | None]:
    """Run one per-group synthesis pass and persist the raw text output.

    Builds the topic lines, formats the prompt, calls the model, and writes
    the result to a per-group text file. On failure, writes a warning and
    returns None for the result.

    Args:
        group:                    Group value (e.g. category name).
        labels_df:                Labeled topics DataFrame.
        groupby_field:            The analysis grouping column.
        group_description:        Prompt-friendly description for this group.
        per_group_instructions:   Instruction block for per-group synthesis.
        client:                   OpenAI client.
        model_name:               Model name (params.yaml → models.synthesis).
        system_prompt:            Synthesis system prompt.
        warnings_path:            JSONL file for recording failures.
        outpath_func:             Callable(subdir, fname) → Path for output files.
        synthesis_top_terms_count: Terms per topic line in the synthesis prompt.
        max_retries:              Maximum retry attempts.
        stage_name:               Stage name for warning records.

    Returns:
        (group, result_text) where result_text is None on failure.
    """
    topic_lines_text = build_topic_lines(
        labels_df, groupby_field, group=group, top_terms_count=synthesis_top_terms_count
    )
    prompt = build_per_group_prompt(
        group=group,
        group_description=group_description,
        topic_lines_text=topic_lines_text,
        per_group_instructions=per_group_instructions,
    )
    try:
        result = _call_with_retry(
            prompt, client=client, model_name=model_name,
            system_prompt=system_prompt, max_retries=max_retries,
        )
    except (_openai.RateLimitError, _openai.APITimeoutError) as e:
        # _call_with_retry re-raises after exhausting retries; catch here so
        # a single group failure takes the graceful warning path rather than
        # halting the notebook.
        append_warning(
            warnings_path, stage_name, "SYNTHESIS_GROUP_FAILED",
            f"Synthesis failed for group '{group}' after retries: {e}",
            context={"group": group, "error": str(e)},
        )
        return group, None
    if result is None:
        append_warning(
            warnings_path, stage_name, "SYNTHESIS_GROUP_FAILED",
            f"Synthesis failed for group '{group}'",
            context={"group": group},
        )
        return group, None

    slug = slugify_group_value(group)
    fpath = outpath_func("analysis", f"llm_synthesis_{slug}.txt")
    with open(fpath, "w", encoding="utf-8") as f:
        f.write(result)
    return group, result


# ── 16. Insight normalization & verification ──────────────────────────────────


def strip_json_fences(text: str | None) -> str:
    """Remove optional ```json code fences from model output before json.loads().

    Args:
        text: Raw model response string, or None.

    Returns:
        Stripped string with fences removed.
    """
    text = (text or "").strip()
    text = re.sub(r"^\s*```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```\s*$", "", text)
    return text.strip()


def normalize_source_topics(
    source_topics: list[Any] | None,
    required_group_values: list[Any],
) -> list[dict[str, Any]]:
    """Normalise model source_topics into [{'group': ..., 'topic_id': int}] records.

    Handles two input formats:
        - Dicts with group/topic_id keys (primary format from structured model output).
        - Pipe-delimited strings of the form "group_value|topic_id".

    Note: bare "Topic N" strings (without a group) are parsed for the topic_id
    but are then discarded because the subsequent group-membership check requires
    a non-empty group value present in required_group_values. They do not survive
    into the output list.

    Entries whose group value is not in required_group_values are always filtered
    out. Output is deduplicated while preserving first-seen order.

    Args:
        source_topics:         Raw source_topics list from model output.
        required_group_values: Allowed group values (from the source topics table).

    Returns:
        List of {group, topic_id} dicts, deduplicated in original order.
    """
    out = []
    for src in source_topics or []:
        group = tid = None
        if isinstance(src, dict):
            group = src.get("group", src.get("group_value", ""))
            tid = src.get("topic_id", src.get("topic", src.get("id", "")))
        elif isinstance(src, str):
            s = src.strip()
            if "|" in s:
                left, right = s.rsplit("|", 1)
                group, tid = left.strip(), right.strip()
            elif re.fullmatch(r"(?i)topic\s+\d+", s):
                tid = re.sub(r"(?i)topic\s+", "", s).strip()

        if group is not None:
            group = str(group).strip()
        if tid is not None:
            tid = str(tid).strip()
        if not tid:
            continue
        try:
            tid = int(float(tid))
        except Exception:
            continue
        if not group or group not in required_group_values:
            continue
        out.append({"group": group, "topic_id": tid})

    seen: set[tuple[str, int]] = set()
    deduped = []
    for item in out:
        key = (item["group"], item["topic_id"])
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped


def normalize_insight(
    insight: dict[str, Any],
    required_group_values: list[Any],
) -> dict[str, Any]:
    """Normalise one insight object, preserving both verified and claimed source_topics.

    Args:
        insight:               Raw insight dict from model JSON output.
        required_group_values: Allowed group values for source_topic filtering.

    Returns:
        Normalised dict with keys: title, what_seeing, why_easy_to_miss,
        source_topics (normalised), source_topics_claimed (raw copy).

    Raises:
        ValueError: If insight is not a dict, or if title is empty.
    """
    if not isinstance(insight, dict):
        raise ValueError(f"Insight must be a dict, got {type(insight)}")
    raw_source_topics = deepcopy(insight.get("source_topics", []))
    source_topics = normalize_source_topics(raw_source_topics, required_group_values)
    normalized = {
        "title": str(insight.get("title", "")).strip(),
        "what_seeing": str(insight.get("what_seeing", "")).strip(),
        "why_easy_to_miss": str(insight.get("why_easy_to_miss", "")).strip(),
        "source_topics": source_topics,
        "source_topics_claimed": raw_source_topics,
    }
    if not normalized["title"]:
        raise ValueError(f"Insight missing title: {insight}")
    return normalized


def project_insight_for_saved_candidates(insight: dict[str, Any]) -> dict[str, Any]:
    """Reformat a saved candidate insight into the normalised pipeline shape.

    Used when loading previously saved candidates (e.g. from a prior run)
    back into the pipeline for re-processing. Converts flexible source_topic
    formats to the canonical group|topic_id string list.

    Args:
        insight: Candidate insight dict from a saved JSON file.

    Returns:
        Dict with keys: title, what_seeing, why_easy_to_miss, source_topics.
    """
    source_topics_out = []
    for src in insight.get("source_topics", []):
        if isinstance(src, dict):
            group = str(src.get("group", "")).strip()
            topic_id = int(float(src.get("topic_id", -1)))
            if group and topic_id >= 0:
                source_topics_out.append(f"{group}|{topic_id}")
        elif isinstance(src, str) and "|" in src:
            source_topics_out.append(src.strip())
    return {
        "title": str(insight.get("title", "")).strip(),
        "what_seeing": str(insight.get("what_seeing", "")).strip(),
        "why_easy_to_miss": str(insight.get("why_easy_to_miss", "")).strip(),
        "source_topics": source_topics_out,
    }


def verify_source_topics(
    insight: dict[str, Any],
    labels_df: pd.DataFrame,
    groupby_field: str,
    required_group_values: list[Any],
    *,
    client: OpenAI,
    model_verify: str,
    system_prompt: str,
    warnings_path: Path,
    max_retries: int = 3,
    stage_name: str = "03_insights_generation",
) -> dict[str, Any]:
    """Verify that each claimed source topic directly supports the insight.

    Sends the insight title, what_seeing text, and each claimed topic's label
    and description to the model. The model returns only the subset of topics
    that genuinely support the claim. Topics dropped by the model are removed
    from insight["source_topics"] in-place.

    Returns the insight unchanged if it has no source_topics or if all API
    attempts fail (warnings are recorded but the pipeline continues).

    Args:
        insight:               Insight dict with source_topics list.
        labels_df:             Labeled topics DataFrame.
        groupby_field:         The analysis grouping column.
        required_group_values: Allowed group values.
        client:                OpenAI client.
        model_verify:          Model name (params.yaml → models.verify).
        system_prompt:         Verification system prompt.
        warnings_path:         JSONL file for recording failures.
        max_retries:           Maximum retry attempts.
        stage_name:            Stage name for warning records.

    Returns:
        The insight dict with source_topics updated to verified-only topics.
    """
    title = insight.get("title", "")
    what_seeing = insight.get("what_seeing", "")
    source_topics = insight.get("source_topics", [])
    warning_id = (str(title).strip() or str(what_seeing).strip())[:120]

    if not source_topics:
        return insight

    topic_lines = []
    for src in source_topics:
        if isinstance(src, dict):
            group = src.get("group", "")
            tid = str(src.get("topic_id", ""))
        elif isinstance(src, str) and "|" in src:
            group, tid = src.rsplit("|", 1)
        else:
            continue
        match = labels_df[
            (labels_df[groupby_field] == group)
            & (labels_df["topic_id"] == int(float(tid)))
        ]
        if not match.empty:
            row = match.iloc[0]
            topic_lines.append(
                f"  {group}|{tid} | label: {row['proposed_label']} | "
                f"description: {row['description']}"
            )

    if not topic_lines:
        return insight

    prompt = f"""
Insight title: {title}

What we're seeing: {what_seeing}

Claimed source topics:
{chr(10).join(topic_lines)}

Return JSON: {{"verified_topics": [{{"group": "...", "topic_id": <int>}}, ...]}}
""".strip()

    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model_verify,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
                response_format={"type": "json_object"},
            )
            result = json.loads(resp.choices[0].message.content)
            original_n = len(source_topics)
            insight["source_topics"] = normalize_source_topics(
                result.get("verified_topics", source_topics), required_group_values
            )
            verified_n = len(insight.get("source_topics", []))
            if verified_n != original_n:
                print(f"Adjusted source_topics: {title[:80]} | {original_n} -> {verified_n}")
            if verified_n == 0 and original_n > 0:
                print(f"WARNING: all source_topics removed: {title[:80]}")
            return insight
        except (_openai.RateLimitError, _openai.APITimeoutError) as e:
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
            else:
                append_warning(warnings_path, stage_name, "VERIFY_API_FAILURE",
                               f"Verification failed for '{warning_id or '[untitled]'}'",
                               context={"title": title or None, "error": str(e)})
                return insight
        except Exception as e:
            append_warning(warnings_path, stage_name, "VERIFY_API_FAILURE",
                           f"Verification failed for '{warning_id or '[untitled]'}'",
                           context={"title": title or None, "error": str(e)})
            return insight


def _verify_insight_list(
    items: list[dict[str, Any]],
    *,
    labels_df: pd.DataFrame,
    groupby_field: str,
    required_group_values: list[Any],
    client: OpenAI,
    model_verify: str,
    system_prompt: str,
    warnings_path: Path,
    min_source_topics_to_verify: int = 1,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """Run verify_source_topics() over a list of insights and collect summary stats.

    Insights with fewer than min_source_topics_to_verify claimed topics are
    passed through without verification (used to skip narrow by-group insights).

    Args:
        items:                       List of insight dicts.
        labels_df:                   Labeled topics DataFrame.
        groupby_field:               The analysis grouping column.
        required_group_values:       Allowed group values.
        client:                      OpenAI client.
        model_verify:                Verification model name.
        system_prompt:               Verification system prompt.
        warnings_path:               JSONL file for recording failures.
        min_source_topics_to_verify: Minimum claimed topic count to trigger
                                     verification. From analysis.verification
                                     .by_group_min_source_topics in params.yaml.

    Returns:
        (verified_items, stats_dict) where stats_dict records counts of
        changed and dropped-to-zero insights.
    """
    verified_items = []
    changed_count = dropped_to_zero_count = topics_before = topics_after = 0

    for insight in items:
        before_n = len(insight.get("source_topics", []))
        topics_before += before_n
        if before_n >= min_source_topics_to_verify:
            verified = verify_source_topics(
                insight, labels_df=labels_df, groupby_field=groupby_field,
                required_group_values=required_group_values, client=client,
                model_verify=model_verify, system_prompt=system_prompt,
                warnings_path=warnings_path,
            )
        else:
            verified = insight
        after_n = len(verified.get("source_topics", []))
        topics_after += after_n
        if after_n != before_n:
            changed_count += 1
        if before_n > 0 and after_n == 0:
            dropped_to_zero_count += 1
        verified_items.append(verified)

    return verified_items, {
        "insight_count": len(items),
        "changed_count": changed_count,
        "dropped_to_zero_count": dropped_to_zero_count,
        "topics_before": topics_before,
        "topics_after": topics_after,
    }


# ── 17. Dedup helpers ─────────────────────────────────────────────────────────


def _norm_text(s: Any) -> str:
    """Normalise a string for token-overlap comparison.

    Lowercases, strips punctuation, and collapses whitespace.
    """
    s = str(s).lower()
    s = re.sub(r"[^a-z0-9\s]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()


def _token_set(s: Any) -> set[str]:
    """Return the set of words in a normalised string."""
    return set(_norm_text(s).split())


def _jaccard(a: Any, b: Any) -> float:
    """Compute Jaccard similarity between two token sets (or iterables).

    Returns 0.0 when either set is empty.
    """
    a, b = set(a), set(b)
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def _topic_list(val: Any) -> list[str]:
    """Parse a source_topics value into a flat list of 'group|topic_id' strings.

    Handles list, JSON string, and comma-separated string formats. Used to
    build the verified_topics_list column for dedup comparisons.
    """
    if isinstance(val, list):
        return [str(x).strip() for x in val if str(x).strip()]
    if pd.isna(val):
        return []
    s = str(val).strip()
    if not s:
        return []
    try:
        parsed = json.loads(s)
        if isinstance(parsed, list):
            return [str(x).strip() for x in parsed if str(x).strip()]
    except Exception:
        pass
    return [x.strip() for x in s.split(",") if x.strip()]


def _pair_kind(a: dict[str, Any], b: dict[str, Any]) -> str | None:
    """Classify an insight pair into a dedup comparison category.

    Categories:
        key_vs_key    — both are cross-group (key_insights) insights.
        bg_vs_same_bg — both are by-group insights from the same category.
        None          — different-bucket by-group pair; never deduped.

    bg_vs_key pairs are intentionally excluded from dedup: cross-group and
    by-group insights can legitimately cover similar ground at different
    levels of specificity.

    Args:
        a: Row dict for the kept insight (must have 'section', 'category_bucket').
        b: Row dict for the candidate insight.

    Returns:
        Category string or None.
    """
    a_key = a["section"] == "key_insights"
    b_key = b["section"] == "key_insights"
    if a_key and b_key:
        return "key_vs_key"
    if not a_key and not b_key:
        return "bg_vs_same_bg" if a["category_bucket"] == b["category_bucket"] else None
    return None  # bg_vs_key — skip dedup


def _screen_pair(
    a: dict[str, Any],
    b: dict[str, Any],
) -> dict[str, Any] | None:
    """Pre-screen an insight pair for potential duplication before the dedupe rules.

    Computes topic, title, and text overlap and returns a metadata dict only
    for pairs that exceed the screening thresholds. Pairs below all thresholds
    are returned as None and never reach the deterministic dedupe logic.

    Screening thresholds (intentionally generous to avoid false negatives):
        topic_overlap  ≥ 0.30
        title_overlap  ≥ 0.45
        text_overlap   ≥ 0.40

    Args:
        a: Row dict for the kept insight.
        b: Row dict for the candidate insight.

    Returns:
        Dict with pair_kind, topic_overlap, title_overlap, text_overlap — or
        None if pair_kind is None or no threshold is met.
    """
    kind = _pair_kind(a, b)
    if kind is None:
        return None

    topic_overlap = _jaccard(a["verified_topics_list"], b["verified_topics_list"])
    title_overlap = _jaccard(a["title_tokens"], b["title_tokens"])
    text_overlap = _jaccard(a["text_tokens"], b["text_tokens"])

    if topic_overlap >= 0.30 or title_overlap >= 0.45 or text_overlap >= 0.40:
        return {
            "pair_kind": kind,
            "topic_overlap": topic_overlap,
            "title_overlap": title_overlap,
            "text_overlap": text_overlap,
        }
    return None


# ── 18. Evidence & support tables ────────────────────────────────────────────


def get_topic_key(groupby_field: str, group: Any, topic_id: Any) -> str:
    """Build the canonical topic_key string used in the bridge table.

    Format: {groupby_field}={group}|topic={topic_id}

    Args:
        groupby_field: The analysis grouping column name.
        group:         Group value.
        topic_id:      Integer topic ID.

    Returns:
        Topic key string.
    """
    return f"{groupby_field}={group}|topic={int(float(topic_id))}"


def iter_candidate_insights(
    data: dict[str, Any],
    output_group_key: str,
) -> Any:
    """Yield every synthesized insight as a flat candidate dict.

    Assigns sequential IDs: KI_001, KI_002, ... for key_insights and
    BG_001, BG_002, ... for by-group insights.

    Args:
        data:             insights_data dict with key_insights and by-group lists.
        output_group_key: Key for the by-group section (e.g. "by_group").

    Yields:
        Dicts with keys: insight_id, section, category_bucket, insight.
    """
    for idx, insight in enumerate(data.get("key_insights", []), start=1):
        yield {
            "insight_id": f"KI_{idx:03d}",
            "section": "key_insights",
            "category_bucket": None,
            "insight": insight,
        }
    group_idx = 1
    for group_value, items in data.get(output_group_key, {}).items():
        for insight in items:
            yield {
                "insight_id": f"BG_{group_idx:03d}",
                "section": output_group_key,
                "category_bucket": group_value,
                "insight": insight,
            }
            group_idx += 1


def _parse_topic_id(val: Any) -> int:
    """Parse a topic ID from either a plain integer or a 'Topic N' string."""
    s = str(val).strip()
    if s.lower().startswith("topic"):
        s = s.split()[-1]
    return int(s)


def build_bridge_lookup(
    project_topic_bridge_df: pd.DataFrame,
) -> dict[str, pd.DataFrame]:
    """Index the project-topic bridge table by topic_key for O(1) lookup.

    Args:
        project_topic_bridge_df: Output of build_project_topic_bridge().

    Returns:
        Dict mapping topic_key → DataFrame of {project_id, weight, topic_share}.
    """
    return {
        topic_key: grp[["project_id", "weight", "topic_share"]].copy()
        for topic_key, grp in project_topic_bridge_df.groupby("topic_key", observed=True)
    }


def build_label_index(
    labels_df: pd.DataFrame,
    groupby_field: str,
    warnings_path: Path | None = None,
    stage_name: str = "03_insights_generation",
) -> dict[tuple[str, int], Any]:
    """Build a (group, topic_id) → label row lookup dict from labels_df.

    Emits a warning for duplicate keys and keeps the first occurrence.

    Args:
        labels_df:     Labeled topics DataFrame from the labeling step.
        groupby_field: The analysis grouping column.
        warnings_path: Optional JSONL file for duplicate-key warnings.
        stage_name:    Stage name for warning records.

    Returns:
        Dict mapping (group_str, topic_id_int) → labels_df row.
    """
    out: dict[tuple[str, int], Any] = {}
    for _, row in labels_df.iterrows():
        key = (str(row[groupby_field]), int(row["topic_id"]))
        if key in out:
            if warnings_path is not None:
                append_warning(
                    warnings_path, stage_name, "DUPLICATE_TOPIC_LABEL_ROW",
                    f"Duplicate labels_df row for {key}; keeping first occurrence",
                    context={"group": key[0], "topic_id": key[1]},
                )
            continue
        out[key] = row
    return out


def summarize_insight_support(
    candidate: dict[str, Any],
    *,
    groupby_field: str,
    run_id: str,
    bridge_lookup: dict[str, pd.DataFrame],
    label_index: dict[tuple[str, int], Any],
    top_project_id_limit: int = 100,
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    """Compute support statistics for one candidate insight.

    For each verified source topic, retrieves its project rows from the bridge
    lookup and its label from the label index. Ranks supporting projects by
    combined topic_share across all source topics.

    Args:
        candidate:            Candidate dict from iter_candidate_insights().
        groupby_field:        The analysis grouping column.
        run_id:               Run ID for provenance.
        bridge_lookup:        From build_bridge_lookup().
        label_index:          From build_label_index().
        top_project_id_limit: Maximum number of project IDs to include in the
                              flat row. Intended to be sourced from
                              params.yaml → output.csv_max_ids_per_insight
                              once that wiring is in place.

    Returns:
        (flat_row, support_rows) where flat_row is one insight summary dict
        and support_rows is a list of per-topic support detail dicts.
    """
    insight = candidate["insight"]
    source_topics = insight.get("source_topics", [])
    claimed_topic_count = len(insight.get("source_topics_claimed", []))

    support_rows: list[dict[str, Any]] = []
    ranking_frames: list[pd.DataFrame] = []

    for src in source_topics:
        if not isinstance(src, dict):
            continue
        group = str(src.get("group", "")).strip()
        topic_id = int(float(src.get("topic_id", -1)))
        topic_key = get_topic_key(groupby_field, group, topic_id)
        topic_rows = bridge_lookup.get(
            topic_key, pd.DataFrame(columns=["project_id", "weight", "topic_share"])
        )
        label_row = label_index.get((group, topic_id))

        if not topic_rows.empty:
            ranking_frames.append(topic_rows[["project_id", "topic_share", "weight"]].copy())

        support_rows.append({
            "run_id": run_id,
            "insight_id": candidate["insight_id"],
            "section": candidate["section"],
            "category_bucket": candidate["category_bucket"],
            "group_by_field": groupby_field,
            "group_value": group,
            "topic_id": topic_id,
            "topic_label": label_row["proposed_label"] if label_row is not None else "[not found]",
            "topic_description": label_row["description"] if label_row is not None else "[not found]",
            "coherence_flag": label_row["coherence_flag"] if label_row is not None else "unknown",
            "supporting_project_count": int(topic_rows["project_id"].nunique()) if not topic_rows.empty else 0,
            "mean_topic_share": float(topic_rows["topic_share"].mean()) if not topic_rows.empty else 0.0,
            "median_topic_share": float(topic_rows["topic_share"].median()) if not topic_rows.empty else 0.0,
        })

    if ranking_frames:
        project_scores = (
            pd.concat(ranking_frames, ignore_index=True)
            .groupby("project_id", as_index=False)
            .agg(
                total_topic_share=("topic_share", "sum"),
                total_weight=("weight", "sum"),
            )
        )
        project_scores["project_id_numeric"] = pd.to_numeric(
            project_scores["project_id"], errors="coerce"
        )
        project_scores = project_scores.sort_values(
            ["total_topic_share", "total_weight", "project_id_numeric", "project_id"],
            ascending=[False, False, True, True],
            na_position="last",
        )
        top_project_ids = project_scores["project_id"].tolist()[:top_project_id_limit]
        supporting_project_count = int(len(project_scores))
    else:
        top_project_ids = []
        supporting_project_count = 0

    verified_topic_count = len(support_rows)
    verification_ratio = (
        verified_topic_count / claimed_topic_count if claimed_topic_count > 0 else 0.0
    )
    mean_topic_share_all = (
        float(np.mean([r["mean_topic_share"] for r in support_rows])) if support_rows else 0.0
    )

    flat_row: dict[str, Any] = {
        "run_id": run_id,
        "insight_id": candidate["insight_id"],
        "section": candidate["section"],
        "category_bucket": candidate["category_bucket"],
        "title": insight.get("title", ""),
        "what_seeing": insight.get("what_seeing", ""),
        "why_easy_to_miss": insight.get("why_easy_to_miss", ""),
        "source_topics_verified": source_topics,
        "source_topics_claimed": insight.get("source_topics_claimed", []),
        "claimed_topic_count": int(claimed_topic_count),
        "verified_topic_count": int(verified_topic_count),
        "verification_ratio": float(verification_ratio),
        "supporting_project_count": int(supporting_project_count),
        "mean_topic_share_all_verified_topics": float(mean_topic_share_all),
        "top_project_ids": top_project_ids,
    }
    return flat_row, support_rows


def build_verified_insight_tables(
    insights_data: dict[str, Any],
    output_group_key: str,
    *,
    groupby_field: str,
    bridge_lookup: dict[str, pd.DataFrame],
    label_index: dict[tuple[str, int], Any],
    run_id: str | None = None,
    top_project_id_limit: int = 100,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Build the flat insights table and per-topic support table for all candidates.

    Iterates every candidate insight via iter_candidate_insights() and calls
    summarize_insight_support() for each one.

    Args:
        insights_data:       Raw insights dict from the synthesis step.
        output_group_key:    By-group key (e.g. "by_group").
        groupby_field:       The analysis grouping column.
        bridge_lookup:       From build_bridge_lookup().
        label_index:         From build_label_index().
        run_id:              Run ID for provenance.
        top_project_id_limit: Maximum project IDs per insight row.

    Returns:
        (insights_flat_df, insight_topic_support_df)
    """
    flat_rows: list[dict[str, Any]] = []
    support_rows: list[dict[str, Any]] = []

    for candidate in iter_candidate_insights(insights_data, output_group_key):
        flat_row, topic_rows = summarize_insight_support(
            candidate,
            groupby_field=groupby_field,
            run_id=run_id,
            bridge_lookup=bridge_lookup,
            label_index=label_index,
            top_project_id_limit=top_project_id_limit,
        )
        flat_rows.append(flat_row)
        support_rows.extend(topic_rows)

    return pd.DataFrame(flat_rows), pd.DataFrame(support_rows)


# ── 19. Packaging & tiering ───────────────────────────────────────────────────


def apply_deterministic_packaging(
    insights_flat_df: pd.DataFrame,
    *,
    output_group_key: str,
    packaging_cfg: dict[str, Any],
) -> dict[str, pd.DataFrame]:
    """Apply quality thresholds to accept insights into the report pool.

    Filters on three thresholds from params.yaml → analysis.packaging:
        min_verified_topic_count
        min_supporting_project_count
        min_mean_topic_share

    The main/appendix split is NOT performed here — that is handled downstream
    by assign_topline_sections_simple(). This function only determines which
    insights clear the quality bar.

    Args:
        insights_flat_df: Output of build_verified_insight_tables().
        output_group_key: By-group key (e.g. "by_group"). Unused directly
                          but kept in signature for caller compatibility.
        packaging_cfg:    Dict from params.yaml → analysis.packaging.

    Returns:
        Dict with a single key "accepted_df" containing the accepted insights,
        sorted by supporting_project_count descending.
    """
    pack_df = insights_flat_df.copy()
    pack_df["verified_topic_count"] = pack_df["verified_topic_count"].fillna(0).astype(int)
    pack_df["claimed_topic_count"] = pack_df["claimed_topic_count"].fillna(0).astype(int)
    pack_df["supporting_project_count"] = pack_df["supporting_project_count"].fillna(0).astype(int)
    pack_df["mean_topic_share_all_verified_topics"] = (
        pack_df["mean_topic_share_all_verified_topics"].fillna(0.0).astype(float)
    )
    if "verification_ratio" not in pack_df.columns:
        pack_df["verification_ratio"] = np.where(
            pack_df["claimed_topic_count"] > 0,
            pack_df["verified_topic_count"] / pack_df["claimed_topic_count"],
            0.0,
        )

    accepted_df = pack_df[
        (pack_df["verified_topic_count"] >= packaging_cfg["min_verified_topic_count"])
        & (pack_df["supporting_project_count"] >= packaging_cfg["min_supporting_project_count"])
        & (pack_df["mean_topic_share_all_verified_topics"] >= packaging_cfg["min_mean_topic_share"])
    ].copy()

    accepted_df = accepted_df.sort_values(
        ["supporting_project_count", "verified_topic_count",
         "mean_topic_share_all_verified_topics", "title"],
        ascending=[False, False, False, True],
    ).reset_index(drop=True)

    return {"accepted_df": accepted_df}


def dedupe_packaged_insights(
    accepted_df: pd.DataFrame,
    *,
    dedupe_cfg: dict[str, Any],
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Remove obvious duplicate insights from the accepted pool.

    Insights are processed in quality-rank order (highest mean_topic_share first).
    Each candidate is compared against all already-kept insights. A candidate is
    dropped if it is a key_vs_key or bg_vs_same_bg pair with a kept insight and
    meets the deterministic duplicate thresholds from params.yaml → analysis.dedupe.

    bg_vs_key pairs are never deduped (cross-group and by-group insights serve
    different purposes even when covering similar territory).

    Args:
        accepted_df:  Output of apply_deterministic_packaging()["accepted_df"].
        dedupe_cfg:   Dict from params.yaml → analysis.dedupe.

    Returns:
        (curated_df, audit_df) where audit_df records every dropped pair
        with overlap scores and reason.
    """
    working_df = accepted_df.copy()
    working_df["verified_topics_list"] = working_df["source_topics_verified"].apply(_topic_list)
    working_df["title_tokens"] = working_df["title"].apply(_token_set)
    working_df["text_tokens"] = (
        working_df["title"].fillna("") + " " + working_df["what_seeing"].fillna("")
    ).apply(_token_set)

    working_df = working_df.sort_values(
        ["mean_topic_share_all_verified_topics", "supporting_project_count", "verification_ratio"],
        ascending=False,
    ).reset_index(drop=True)

    kept_rows: list[dict[str, Any]] = []
    audit_rows: list[dict[str, Any]] = []

    for _, row in working_df.iterrows():
        row_dict = row.to_dict()
        drop_current = False

        for kept in kept_rows:
            screen_meta = _screen_pair(row_dict, kept)
            if screen_meta is None or screen_meta["pair_kind"] == "bg_vs_key":
                continue

            topic_overlap = screen_meta["topic_overlap"]
            title_overlap = screen_meta["title_overlap"]
            text_overlap = screen_meta["text_overlap"]

            is_obvious_duplicate = (
                (topic_overlap >= dedupe_cfg["topic_overlap_high"])
                or (
                    topic_overlap >= dedupe_cfg["topic_overlap_min"]
                    and title_overlap >= dedupe_cfg["title_overlap_min"]
                    and text_overlap >= dedupe_cfg["text_overlap_min"]
                )
            )
            if is_obvious_duplicate:
                drop_current = True
                audit_rows.append({
                    "dropped_insight_id": row_dict["insight_id"],
                    "matched_kept_insight_id": kept["insight_id"],
                    "pair_kind": screen_meta["pair_kind"],
                    "topic_overlap": topic_overlap,
                    "title_overlap": title_overlap,
                    "text_overlap": text_overlap,
                    "reason": "deterministic_obvious_duplicate",
                })
                break

        if not drop_current:
            kept_rows.append(row_dict)

    return pd.DataFrame(kept_rows).copy(), pd.DataFrame(audit_rows).copy()


def assign_topline_sections_simple(
    curated_df: pd.DataFrame,
    *,
    output_group_key: str,
    main_cross_limit: int,
    main_min_verification_ratio: float = 0.0,
) -> pd.DataFrame:
    """Assign report_section and report_order to each curated insight.

    Main section:
        - Top N cross-category insights (by quality rank), up to main_cross_limit,
          optionally filtered by main_min_verification_ratio.
        - Top 1 by-group insight per category (by quality rank).
    Appendix:
        - All remaining accepted insights.

    Args:
        curated_df:                    Output of dedupe_packaged_insights()[0].
        output_group_key:              By-group section key (e.g. "by_group").
        main_cross_limit:              Max cross-category insights in the main section.
                                       From params.yaml → analysis.packaging.main_cross_limit.
        main_min_verification_ratio:   Minimum verification_ratio for main-section
                                       eligibility. From params.yaml →
                                       analysis.packaging.main_min_verification_ratio.
                                       Default 0.0 applies no filter (backward compatible).

    Returns:
        DataFrame with report_section ("main_cross", "main_by_group",
        "appendix_cross", "appendix_by_group") and report_order columns added.
    """
    rank_cols = [
        "supporting_project_count", "verified_topic_count",
        "mean_topic_share_all_verified_topics", "title",
    ]
    rank_asc = [False, False, False, True]
    working_df = curated_df.copy()

    # Cross-category main section — apply verification ratio filter if configured.
    cross_pool = working_df[working_df["section"] == "key_insights"]
    if main_min_verification_ratio > 0.0 and "verification_ratio" in cross_pool.columns:
        cross_pool = cross_pool[cross_pool["verification_ratio"] >= main_min_verification_ratio]

    main_cross_df = (
        cross_pool.sort_values(rank_cols, ascending=rank_asc)
        .head(main_cross_limit)
        .copy()
    )
    main_cross_df["report_section"] = "main_cross"
    main_cross_df["report_order"] = range(1, len(main_cross_df) + 1)

    # By-group main section — top 1 per category.
    by_group_pool = working_df[working_df["section"] == output_group_key]
    if main_min_verification_ratio > 0.0 and "verification_ratio" in by_group_pool.columns:
        by_group_pool = by_group_pool[by_group_pool["verification_ratio"] >= main_min_verification_ratio]

    main_by_group_df = (
        by_group_pool.sort_values(
            ["category_bucket", "supporting_project_count", "verified_topic_count",
             "mean_topic_share_all_verified_topics", "title"],
            ascending=[True, False, False, False, True],
        )
        .groupby("category_bucket", as_index=False, sort=True)
        .head(1)
        .copy()
    )
    main_by_group_df["report_section"] = "main_by_group"
    main_by_group_df["report_order"] = range(1, len(main_by_group_df) + 1)

    # Appendix — everything not in main.
    main_ids = set(main_cross_df["insight_id"]) | set(main_by_group_df["insight_id"])
    remainder_df = working_df[~working_df["insight_id"].isin(main_ids)].copy()

    appendix_cross_df = (
        remainder_df[remainder_df["section"] == "key_insights"]
        .sort_values(rank_cols, ascending=rank_asc)
        .copy()
    )
    appendix_cross_df["report_section"] = "appendix_cross"
    appendix_cross_df["report_order"] = range(1, len(appendix_cross_df) + 1)

    appendix_by_group_df = (
        remainder_df[remainder_df["section"] == output_group_key]
        .sort_values(
            ["category_bucket", "supporting_project_count", "verified_topic_count",
             "mean_topic_share_all_verified_topics", "title"],
            ascending=[True, False, False, False, True],
        )
        .copy()
    )
    appendix_by_group_df["report_section"] = "appendix_by_group"
    appendix_by_group_df["report_order"] = (
        appendix_by_group_df.groupby("category_bucket", sort=True).cumcount() + 1
    )

    return pd.concat(
        [main_cross_df, main_by_group_df, appendix_cross_df, appendix_by_group_df],
        ignore_index=True,
        sort=False,
    )


def build_structured_from_curated(
    curated_df: pd.DataFrame,
    *,
    output_group_key: str,
) -> dict[str, Any]:
    """Convert the curated insights DataFrame into the nested JSON structure for the report.

    Produces the insights_structured.json format consumed by build_packaged_report_docx().

    Args:
        curated_df:       Output of assign_topline_sections_simple().
        output_group_key: By-group section key (e.g. "by_group").

    Returns:
        Dict with keys: "key_insights" (list) and output_group_key (dict of
        category → list of insight dicts).
    """
    structured: dict[str, Any] = {"key_insights": [], output_group_key: {}}

    for _, row in curated_df.sort_values(["report_section", "report_order", "title"]).iterrows():
        looker_url = row.get("looker_url", "")
        if not isinstance(looker_url, str) or pd.isna(looker_url):
            looker_url = ""

        top_project_ids = row.get("top_project_ids", [])
        if not isinstance(top_project_ids, list):
            top_project_ids = []

        item: dict[str, Any] = {
            "insight_id": row["insight_id"],
            "title": row["title"],
            "what_seeing": row["what_seeing"],
            "why_easy_to_miss": row["why_easy_to_miss"],
            "source_topics": row["source_topics_verified"],
            "supporting_project_count": int(row["supporting_project_count"]),
            "verified_topic_count": int(row["verified_topic_count"]),
            "verification_ratio": float(row["verification_ratio"]),
            "mean_topic_share_all_verified_topics": float(row["mean_topic_share_all_verified_topics"]),
            "top_project_ids": top_project_ids,
            "looker_url": looker_url,
            "report_section": row["report_section"],
            "report_order": int(row["report_order"]),
            "warnings": [],
        }

        if row["section"] == "key_insights":
            structured["key_insights"].append(item)
        else:
            structured[output_group_key].setdefault(row["category_bucket"], []).append(item)

    return structured


# ── 20. DOCX report helpers ───────────────────────────────────────────────────


def build_looker_project_url(
    *,
    base_url: str,
    project_ids: list[Any],
    filter_field: str,
    fields: list[str] | str,
    limit: int = 500,
    max_ids: int = 100,
) -> str:
    """Build a Looker Explore URL pre-filtered to a capped list of project IDs.

    Deduplicates IDs (preserving rank order), caps at max_ids, and URL-encodes
    all parameters. Returns "" when no project IDs remain after normalisation.

    Args:
        base_url:     Looker Explore base URL (params.yaml → output.looker_base_url).
        project_ids:  Ranked list of project IDs to include in the filter.
        filter_field: Looker field to filter on (params.yaml → output.looker_filter_field).
        fields:       List of fields to include, or a comma-joined string.
        limit:        Looker row limit (params.yaml → output.looker_limit).
        max_ids:      Maximum IDs in the URL filter (params.yaml → output.looker_id_limit).

    Returns:
        Full URL string, or "" if no valid project IDs are provided.
    """
    def _normalize(value: Any) -> str:
        if pd.isna(value):
            return ""
        try:
            numeric = float(value)
            if numeric.is_integer():
                return str(int(numeric))
        except (TypeError, ValueError):
            pass
        return str(value).strip()

    normalized = list(dict.fromkeys(pid for v in project_ids if (pid := _normalize(v))))[:max_ids]
    if not normalized:
        return ""

    if isinstance(fields, (list, tuple)):
        fields_param = ",".join(str(f).strip() for f in fields if str(f).strip())
    else:
        fields_param = str(fields).strip()
    if not fields_param:
        raise ValueError("build_looker_project_url requires at least one field")

    params = {
        "fields": fields_param,
        f"f[{filter_field}]": ",".join(normalized),
        "limit": str(limit),
    }
    return f"{base_url}?{urlencode(params)}"


def add_hyperlink(paragraph: Any, text: str, url: str) -> Any:
    """Add an external hyperlink run to a python-docx paragraph.

    python-docx does not natively support hyperlinks, so this function
    manipulates the underlying XML directly.

    Args:
        paragraph: python-docx Paragraph object.
        text:      Visible link text.
        url:       Target URL.

    Returns:
        The created hyperlink XML element.
    """
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_heading(doc: Any, text: str, level: int) -> Any:
    """Add a black Arial heading to a python-docx Document.

    python-docx headings default to the theme colour; this function overrides
    that to ensure consistent black text in the report.

    Args:
        doc:   python-docx Document object.
        text:  Heading text.
        level: Heading level (1–4).

    Returns:
        The created Paragraph object.
    """
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p.runs[0].font.name = "Arial"
    return p


def add_insight_meta_line(
    doc: Any,
    insight: dict[str, Any],
    font_size_pt: float = 8.5,
) -> None:
    """Add a compact italic meta line above an insight body in the DOCX.

    Displays supporting project count, verified source topic count, and
    average topic fit score. Controlled by params.yaml →
    output.report_include_meta_line.

    Args:
        doc:          python-docx Document object.
        insight:      Insight dict containing support statistics.
        font_size_pt: Font size for the meta line. Default 8.5pt.
    """
    mean_topic_fit = insight.get("mean_topic_share_all_verified_topics")
    fit_str = (
        f"{round(mean_topic_fit * 100):.0f}%"
        if isinstance(mean_topic_fit, (int, float))
        else "—"
    )
    text = (
        f"Supporting projects: {insight.get('supporting_project_count', '—')}  |  "
        f"Verified source topics: {insight.get('verified_topic_count', '—')}  |  "
        f"Average topic fit: {fit_str}"
    )
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(font_size_pt)
    p.paragraph_format.space_after = Pt(2)


def add_insight(
    doc: Any,
    insight: dict[str, Any],
    *,
    include_looker_link: bool = True,
    looker_link_text: str = "Project essays for this insight",
) -> None:
    """Add one accepted insight block to the DOCX report.

    Renders: bold title, "What we're seeing:" body, "Why this is easy to miss:"
    body, and (optionally) a Looker Explore hyperlink.

    Args:
        doc:                python-docx Document object.
        insight:            Insight dict from the structured output.
        include_looker_link: When True, appends a Looker link if looker_url
                             is present. Controlled by params.yaml →
                             output.report_include_looker_link.
        looker_link_text:   Visible hyperlink text. From params.yaml →
                            output.report_looker_link_text.
    """
    p = doc.add_paragraph()
    run = p.add_run(insight.get("title", ""))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3E, 0x00, 0xC9)

    for label, key in [
        ("What we're seeing:", "what_seeing"),
        ("Why this is easy to miss:", "why_easy_to_miss"),
    ]:
        p = doc.add_paragraph()
        label_run = p.add_run(label + "  ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        body_run = p.add_run(insight.get(key, ""))
        body_run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    if include_looker_link and insight.get("looker_url"):
        p = doc.add_paragraph()
        label_run = p.add_run("Explore supporting projects: ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        add_hyperlink(p, looker_link_text, insight["looker_url"])
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


def build_packaged_report_docx(
    *,
    structured: dict[str, Any],
    output_path: Any,
    output_group_key: str = "by_group",
    report_cfg: dict[str, Any],
    project_count: int | None = None,
    run_id: str | None = None,
    normal_font_name: str = "Arial",
    normal_font_size_pt: float = 10.0,
    margin_inches: float = 1.0,
) -> None:
    """Build and save the final packaged DOCX report.

    Renders main and appendix sections, each containing cross-category and
    by-group subsections. Section labels, meta lines, and Looker links are
    all controlled by params.yaml → output.report_* keys passed through
    report_cfg.

    Args:
        structured:        Nested insights dict from build_structured_from_curated().
        output_path:       Destination file path for the DOCX.
        output_group_key:  By-group section key (e.g. "by_group").
        report_cfg:        Dict from params.yaml → output block.
        project_count:     Optional project count for the summary line.
        run_id:            Optional run ID for the summary line.
        normal_font_name:  Default body font. Default "Arial".
        normal_font_size_pt: Default body font size. Default 10pt.
        margin_inches:     Page margins. Default 1 inch.
    """
    title = report_cfg.get("report_title", "Trend Report")
    main_label = report_cfg.get("report_main_section_label", "Main Insights")
    appendix_label = report_cfg.get("report_appendix_section_label", "Appendix")
    main_cross_label = report_cfg.get("report_main_cross_label", "Cross-Category Similarities")
    main_by_group_label = report_cfg.get("report_main_by_group_label", "Group-Specific Findings")
    appendix_cross_label = report_cfg.get("report_appendix_cross_label", "Additional Cross-Category Insights")
    appendix_by_group_label = report_cfg.get("report_appendix_by_group_label", "Additional Group-Specific Findings")
    incl_meta = report_cfg.get("report_include_meta_line", True)
    incl_summary = report_cfg.get("report_include_signal_summary", True)
    incl_looker = report_cfg.get("report_include_looker_link", True)
    looker_link_text = report_cfg.get("report_looker_link_text", "Project essays for this insight")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(margin_inches)
    style = doc.styles["Normal"]
    style.font.name = normal_font_name
    style.font.size = Pt(normal_font_size_pt)

    # Report title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x3E, 0x00, 0xC9)
    p.paragraph_format.space_after = Pt(6)

    # Optional summary line
    if incl_summary:
        all_items = structured.get("key_insights", []) + [
            i for items in structured.get(output_group_key, {}).values() for i in items
        ]
        parts = []
        if project_count is not None:
            parts.append(f"Projects in study: {project_count:,}")
        parts.append(f"Accepted insights: {len(all_items)}")
        if run_id is not None:
            parts.append(f"Run ID: {'_'.join(str(run_id).split('_')[:2])}")
        p = doc.add_paragraph()
        run = p.add_run("  |  ".join(parts))
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()

    def _items_for_section(section_name: str):
        """Return (cross_list, by_group_dict) for the given report_section value."""
        cross = [i for i in structured.get("key_insights", []) if i.get("report_section") == section_name]
        by_group = {
            g: [i for i in items if i.get("report_section") == section_name]
            for g, items in structured.get(output_group_key, {}).items()
        }
        return cross, by_group

    def _render_section(section_title, cross_title, by_group_title, cross_key, by_group_key):
        """Render one report section (main or appendix) with cross and by-group subsections."""
        cross_main, bg_main = _items_for_section(cross_key)
        cross_app, bg_app = _items_for_section(by_group_key)
        cross = cross_main + cross_app
        merged: dict[str, list] = {}
        for g, items in {**bg_main, **bg_app}.items():
            merged[g] = merged.get(g, []) + items

        if not cross and not any(merged.values()):
            return

        p = doc.add_paragraph()
        run = p.add_run(section_title)
        run.bold = True
        run.underline = True
        run.font.name = "Arial"
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        if cross:
            p = doc.add_paragraph()
            run = p.add_run(cross_title)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(10)
            for insight in sorted(cross, key=lambda x: (x.get("report_order", 9999), x.get("title", ""))):
                if incl_meta:
                    add_insight_meta_line(doc, insight)
                add_insight(doc, insight, include_looker_link=incl_looker, looker_link_text=looker_link_text)

        if any(merged.values()):
            p = doc.add_paragraph()
            run = p.add_run(by_group_title)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(10)
            for group, items in sorted(merged.items()):
                if not items:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(str(group))
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                for insight in sorted(items, key=lambda x: (x.get("report_order", 9999), x.get("title", ""))):
                    if incl_meta:
                        add_insight_meta_line(doc, insight)
                    add_insight(doc, insight, include_looker_link=incl_looker, looker_link_text=looker_link_text)

    _render_section(main_label, main_cross_label, main_by_group_label, "main_cross", "main_by_group")
    _render_section(appendix_label, appendix_cross_label, appendix_by_group_label, "appendix_cross", "appendix_by_group")
    doc.save(output_path)


# ── 21. __all__ ───────────────────────────────────────────────────────────────

__all__ = [
    # Constants
    "ROOT",
    "PIPELINE_VERSION",
    # Config & I/O
    "resolve_params_path",
    "load_cfg",
    "write_json",
    "compute_md5",
    "artifact_meta",
    "build_output_path",
    "build_run_output_path",
    "outpath",           # deprecated alias — prefer build_output_path
    "get_run_date",
    # Run identity & filter helpers
    "canonicalize_filter_spec",
    "get_filter_fields_key",
    "get_run_id",
    "validate_filter_spec",
    "apply_filters",
    # Stage & pipeline manifests
    "start_stage_manifest",
    "finalize_stage_manifest",
    "build_pipeline_manifest",
    # Warning file helpers
    "ensure_warning_file",
    "append_warning",
    "get_warning_count",
    # LLM client
    "get_llm_client",
    # Ingest
    "ingest",
    # Token helpers
    "tokens_to_str",
    "flat_freq",
    "token_doc_freq",
    "normalize_tokens",
    # Consolidation helpers
    "build_consolidation_candidates",
    # Analysis helpers
    "make_vec",
    "add_bin",
    "group_key",
    "build_project_topic_bridge",
    "slugify_group_value",
    "load_essay_snippet_lookup",
    # Quality
    "HARD_STOPWORDS",
    "quality_report",
    # TF-IDF & NMF helpers
    "cat_tfidf_slice",
    "choose_n_components",
    "nmf_one",
    # Enrichment helpers
    "gate_cluster",
    "classify_batch",
    "inject_tokens",
    # Topic labeling helpers
    "_norm_group_value",
    "_safe_topic_id",
    "build_input",
    "_make_label_error",
    "_label_with_retry",
    # Synthesis helpers
    "clean_label",
    "build_topic_lines",
    "build_per_group_prompt",
    "_call_with_retry",
    "synthesize_one_group",
    # Insight normalization & verification
    "strip_json_fences",
    "normalize_source_topics",
    "normalize_insight",
    "project_insight_for_saved_candidates",
    "verify_source_topics",
    "_verify_insight_list",
    # Dedup helpers
    "_norm_text",
    "_token_set",
    "_jaccard",
    "_topic_list",
    "_pair_kind",
    "_screen_pair",
    # Evidence & support tables
    "get_topic_key",
    "iter_candidate_insights",
    "_parse_topic_id",
    "build_bridge_lookup",
    "build_label_index",
    "summarize_insight_support",
    "build_verified_insight_tables",
    # Packaging & tiering
    "apply_deterministic_packaging",
    "dedupe_packaged_insights",
    "assign_topline_sections_simple",
    "build_structured_from_curated",
    # DOCX report helpers
    "add_hyperlink",
    "add_heading",
    "add_insight_meta_line",
    "add_insight",
    "build_packaged_report_docx",
    "build_looker_project_url",
]
